
"""
Lab Test Mapping Console — FastAPI backend
Reuses existing processor.py / learner.py; chat powered by Claude (Anthropic).

Run:  uvicorn server:app --reload --port 8007
"""
from __future__ import annotations

import json
import os
import re
import sys
import uuid
from pathlib import Path
from typing import Any

# Load .env from this directory
_env_file = Path(__file__).parent / ".env"
if _env_file.exists():
    for _line in _env_file.read_text().splitlines():
        _line = _line.strip()
        if _line and not _line.startswith("#") and "=" in _line:
            _k, _v = _line.split("=", 1)
            os.environ[_k.strip()] = _v.strip()

import sqlite3
import httpx
import anthropic as anthropic_sdk
from rapidfuzz import process as fuzz_process, fuzz
from fastapi import BackgroundTasks, FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# ── Paths ──────────────────────────────────────────────────────────────────────
PROJECT_ROOT  = Path(__file__).parent.parent
FRONTEND_DIR  = Path(__file__).parent
BACKEND_DIR   = PROJECT_ROOT / "webapp" / "backend"

# Reuse existing processor / learner
sys.path.insert(0, str(BACKEND_DIR))
from processor import parse_file, run_match_script, generate_output_excel  # noqa: E402
from learner import apply_learnings                                          # noqa: E402

# ── App ────────────────────────────────────────────────────────────────────────
print(f"[STARTUP] Loading server.py from: {__file__}", flush=True)
app = FastAPI(title="Lab Test Mapping Console", version="2.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# Claude — chat + parsing + semantic recovery
claude = anthropic_sdk.Anthropic(
    api_key=os.environ.get("ANTHROPIC_API_KEY"),
    http_client=httpx.Client(verify=False),
)

# ── Lab-type inference ────────────────────────────────────────────────────────
_RAD_PAT = re.compile(
    r'\b(xray|x[\s\-]ray|mri\b|ct\b|hrct\b|ctscan|usg\b|ultrasound|sonograph|doppler|mammo|'
    r'dexa\b|pet\b|spect\b|nuclear|scintigraph|angiograph|fluoros|barium|ivu\b|hsg\b|rgu\b|mcu\b|'
    r'echo\b|echocardiograph|holter|tmt\b|eeg\b|emg\b|ncv\b|audiometr|fnac\b|'
    r'colonoscop|endoscop|bronchoscop|laparoscop|cystoscop|sigmoidoscop)',
    re.IGNORECASE
)

def _infer_lab_type(name: str) -> str:
    return "Radiology" if _RAD_PAT.search(name) else "Pathology"

_PRICE_COLS_SET = {
    "price", "rate", "mrp", "amount", "charges", "cost",
    "our rate", "our price", "net rate", "list price",
    "price (rs)", "rate (rs)", "rate(rs)", "charges (rs)", "price(rs)",
    "rs", "amount (rs)", "amount(rs)", "net amount", "net price",
}
_DEPT_COLS_SET = {
    "department", "dept", "category", "section", "division", "lab type", "lab",
}

_PRICE_KW_RE = re.compile(
    r'\b(price|rate|mrp|amount|charges|cost)\b', re.IGNORECASE
)


def _is_price_col(header: str) -> bool:
    """True if header looks like a price column — handles ₹, Rs, (Rs.), etc."""
    h = str(header).strip().lower()
    return h in _PRICE_COLS_SET or bool(_PRICE_KW_RE.search(h))


_NAME_HDR_RE  = re.compile(r'name|test|invest|item|descr|param|service|proc|particular', re.I)
_PRICE_HDR_RE = re.compile(r'price|rate|mrp|amount|charge|fee|cost|rs\.?\b|rupee', re.I)
_SKIP_HDR_RE  = re.compile(
    r'^(s\.?\s*[il][\./]?\s*no\.?|sr\.?\s*no\.?|s\.?\s*no\.?|serial|sno|no\.?|#|sn|sl\.?\s*no\.?'
    r'|type|specimen|tat|collect|home|gender|age|loinc|remark|note|status|qty|code|dept|department)$',
    re.I,
)


def _find_name_price_cols(
    headers: list,
    rows: list[list],
) -> tuple[int | None, int | None]:
    """
    Universal content+keyword scorer — finds the test-name column and price column
    in any tabular data without relying on hardcoded header names.

    Strategy:
    - Header keywords give an initial bias (+/-)
    - Column content (text richness vs numeric) refines the score
    - The column with the highest name-score is the name column
    - The column with the highest price-score (numeric, plausible range) is the price column

    Returns (name_col_idx, price_col_idx) — either may be None if not found.
    """
    n = len(headers)
    if n == 0:
        return None, None

    name_s  = [0.0] * n
    price_s = [0.0] * n

    # ── Header signals ────────────────────────────────────────────────────────
    for ci, h in enumerate(headers):
        h_ = str(h or "").strip()
        hl = h_.lower()
        if not hl:
            continue
        if _SKIP_HDR_RE.match(hl):
            name_s[ci]  -= 8.0
            price_s[ci] -= 8.0
            continue
        if _NAME_HDR_RE.search(hl):
            name_s[ci] += 4.0
        if _PRICE_HDR_RE.search(hl) or _is_price_col(hl):
            price_s[ci] += 5.0

    # ── Content signals (sample up to 40 rows) ────────────────────────────────
    for row in rows[:40]:
        cells = list(row)
        for ci in range(min(n, len(cells))):
            v = str(cells[ci] or "").strip()
            if not v or v.lower() in ("nan", "none", "", "-", "n/a"):
                continue
            # Try numeric
            try:
                num = float(v.replace(",", "").replace("₹", "").replace("Rs", "").strip())
                if 10 <= num <= 1_500_000:
                    price_s[ci] += 0.5
                else:
                    price_s[ci] -= 0.05
                name_s[ci] -= 0.15   # numeric → not a test name
            except ValueError:
                price_s[ci] -= 0.15  # non-numeric → not a price
                chars = len(v)
                if 3 <= chars <= 120:
                    name_s[ci] += 0.3
                elif chars > 120:
                    name_s[ci] -= 0.2  # paragraph text, not a test name

    # ── Pick best columns ─────────────────────────────────────────────────────
    best_name, best_ns = None, 0.3   # require at least small positive score
    for ci in range(n):
        if name_s[ci] > best_ns and name_s[ci] > price_s[ci]:
            best_ns, best_name = name_s[ci], ci

    best_price, best_ps = None, 1.0  # require meaningful numeric content
    for ci in range(n):
        if ci == best_name and price_s[ci] < 10:  # only skip if not a clear price header
            continue
        if price_s[ci] > best_ps:
            best_ps, best_price = price_s[ci], ci

    return best_name, best_price


# ── In-memory state ────────────────────────────────────────────────────────────
jobs: dict[str, dict[str, Any]] = {}          # job_id → job data
sessions: dict[str, list[dict]] = {}          # job_id → message history

# ── Disk cache — jobs + chat survive server restarts ──────────────────────────
CACHE_DIR = FRONTEND_DIR / ".cache"
CACHE_DIR.mkdir(exist_ok=True)


def _save_job(job_id: str) -> None:
    """Write job mappings + stats to disk."""
    job = jobs.get(job_id)
    if not job:
        return
    payload = {
        "job_id":      job_id,
        "filename":    job.get("filename", ""),
        "status":      "done",
        "stats":       job.get("stats", {}),
        "mappings":    job.get("mappings", []),
        "skipped":     job.get("skipped", []),
        "output_path": job.get("output_path"),
    }
    try:
        (CACHE_DIR / f"{job_id}.json").write_text(
            json.dumps(payload, ensure_ascii=False), encoding="utf-8"
        )
    except Exception as exc:
        print(f"WARNING: job cache write failed ({exc})", file=sys.stderr)


def _save_chat(job_id: str) -> None:
    """Write chat history to disk after each exchange."""
    history = sessions.get(job_id)
    if history is None:
        return
    try:
        (CACHE_DIR / f"{job_id}_chat.json").write_text(
            json.dumps(history, ensure_ascii=False), encoding="utf-8"
        )
    except Exception as exc:
        print(f"WARNING: chat cache write failed ({exc})", file=sys.stderr)


def _load_caches() -> None:
    """On startup: reload all completed jobs and chat histories from disk."""
    for f in sorted(CACHE_DIR.glob("*.json")):
        if f.name.endswith("_chat.json"):
            continue
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            jid  = data.get("job_id") or f.stem
            if data.get("status") == "done" and jid not in jobs:
                jobs[jid] = {
                    "status":       "done",
                    "progress":     100,
                    "current_step": "Done",
                    "filename":     data.get("filename", ""),
                    "file_path":    "",
                    "matched":      [],
                    "unmatched":    [],
                    "skipped":      data.get("skipped", []),
                    "mappings":     data.get("mappings", []),
                    "stats":        data.get("stats", {}),
                    "output_path":  data.get("output_path"),
                    "error":        None,
                }
        except Exception as exc:
            print(f"WARNING: job cache load failed ({f.name}): {exc}", file=sys.stderr)

    for f in CACHE_DIR.glob("*_chat.json"):
        try:
            jid = f.name[: -len("_chat.json")]
            if jid not in sessions:
                sessions[jid] = json.loads(f.read_text(encoding="utf-8"))
        except Exception as exc:
            print(f"WARNING: chat cache load failed ({f.name}): {exc}", file=sys.stderr)


_load_caches()   # run once at import time

# ── Helpers ────────────────────────────────────────────────────────────────────

def _build_system(job_id: str | None, client_mappings: list[dict] | None = None) -> str:
    """Build system prompt = CLAUDE.md + assistant persona + mapping snapshot.

    client_mappings: rows sent by the frontend (always fresh, survives server restart).
    Falls back to in-memory jobs dict if not provided.
    """
    # Deliberately exclude CLAUDE.md — it references master_file/pipeline internals
    # which bias GPT-4o toward catalogue lookups instead of pure medical judgment.
    claude_md = ""

    persona = """
## Medical Assistant & Mapping Console

You are a knowledgeable medical assistant embedded in a Lab Test Mapping Console.
You have TWO roles — switch between them based on what the user is asking:

### Role 1 — Medical Expert
For ANY medical question (symptoms, diseases, lab tests, radiology, drugs, physiology, clinical guidelines, normal ranges, what a test diagnoses, etc.) — answer directly and thoroughly using your own medical knowledge. Do NOT deflect, do NOT say "consult a doctor", do NOT say it's outside your scope. Give clear, accurate, clinically useful answers. You are the expert.

### Role 2 — Mapping Console Assistant
For mapping-related tasks (review, correct, approve, revalidate lab test name mappings) — follow the rules below.

### OUTPUT FORMAT RULE — applies to every mapping-related response
Whenever you show mapping data, analysis, comparisons, or validation results — ALWAYS use a markdown table. Never use bullet lists for mapping rows. No exceptions.

For validation/revalidation responses, the table MUST have these columns:
| Raw Name | Current Standard Name | Confidence | Verdict | Issue | Proposed Fix |

For simple update confirmations, use:
| Raw Name | Updated Standard Name | Change Made |

### MAPPING RULES — follow these without exception
1. You NEVER ask for row IDs or UUIDs. The system matches rows by raw_name automatically.
2. When a user says "update X to Y" or "change X to Y" — emit the action block IN THE SAME RESPONSE.
3. Identify rows ONLY by their `raw_name` string. Copy it exactly from the mapping data below.
4. If the user refers to a test name (even approximately), find the closest match in raw_name and act on it immediately.
5. NEVER promise a change without including the action block in that same response.
6. You CAN update ANY row — matched, unmatched, or any status.
7. Every row in the snapshot is editable. Never say a row is out of scope.
8. NEVER end with "let me know if you'd like me to update" or "if you would like me to implement" — always state your verdict and ask "Shall I apply these X corrections?" directly.

### Action block format (mapping changes only)
```action
{
  "type": "update_mappings",
  "summary": "One-line summary of what changed",
  "changes": [
    {"raw_name": "<exact raw_name from mapping data below>", "field": "catalogue_name", "value": "New Standard Name"},
    {"raw_name": "<exact raw_name from mapping data below>", "field": "status", "value": "matched"}
  ]
}
```
Supported fields — ALL are fully implemented and writable:
- `field = "catalogue_name"` → sets the standard test name. NEVER use this to change department. For skipped rows, this stores the name but does NOT change status — the user must confirm manually.
- `field = "status"` → "matched" | "unmatched" | "rejected". NEVER set a skipped row's status to "matched" directly — leave status changes for the user to confirm via the UI.
- `field = "price"` → price as a numeric string e.g. "1500"
- `field = "lab_type"` → ONLY "Radiology" or "Pathology". THIS is the ONLY field for department changes.
- `field = "provider_item_name"` → provider's internal name / slug
- `field = "loinc_id"` → LOINC code string
- `field = "lab_requirement"` → services offered e.g. "Blood", "Urine"
- `field = "entity_type"` → catalogue type e.g. "Pathology", "Radiology"
- `field = "mrp"` → partner MRP as string e.g. "500"
- `field = "discounted_price"` → discounted price as string
- `field = "display_mrp"` → display MRP as string
- `field = "collection_type"` → "Lab" | "Home Collection" | "Hybrid"
- `field = "home_collection"` → "Yes" or "No"
- `field = "fasting_required"` → "Yes" or "No"
- `field = "fasting_hours"` → fasting duration e.g. "8-12 hours fasting is mandatory"
- `field = "age_range"` → e.g. "All ages" or "18-65"
- `field = "gender"` → "Male" | "Female" | "Both"
- `field = "minimum_patient"` → minimum sample count as string e.g. "1"
- `field = "report_tat"` → TAT in hours as string e.g. "24"
- `field = "alias"` → comma-separated aliases
- `field = "tags"` → comma-separated tags
- `field = "prescription_required"` → "Yes" or "No"
- `field = "precautions"` → precautions/instructions text
- `field = "description"` → overview/description text

**CRITICAL — department changes MUST use `lab_type`:**
When user says "change dept to Pathology" or "change department from Radiology to Pathology":
```action
{
  "type": "update_mappings",
  "summary": "Changed department to Pathology",
  "changes": [
    {"raw_name": "CT Head - Without Contrast", "field": "lab_type", "value": "Pathology"}
  ]
}
```
NEVER use `catalogue_name` or `status` to encode a department change. `lab_type` is the ONLY correct field.
Copy `raw_name` exactly as it appears in the mapping data.

### Revalidation instructions
When the user asks to "revalidate", "validate low confidence", "verify accuracy", "check mappings", or anything similar:

**You are a senior radiologist/pathologist reviewing this mapping sheet. Use ONLY your own medical knowledge — do NOT reference any master file, catalogue list, or external source.**

**STRICT TWO-STEP PROCESS:**
- **Step 1 (Analysis):** Evaluate every row. State your verdict. Propose fixes. Ask for approval. Do NOT emit an action block.
- **Step 2 (Apply):** Only after the user explicitly says yes/approve/go ahead — emit the action block. Never before.

**⚠️ In Step 1, you MUST NOT include any ```action``` block. It will be applied immediately without user review. Only include the action block in Step 2.**

Use the validation table format defined in the OUTPUT FORMAT RULE above (Raw Name | Current Standard Name | Confidence | Verdict | Issue | Proposed Fix).
- Verdict: ✅ Correct or ❌ Wrong
- Issue: blank for correct rows; one-phrase reason for wrong ones (e.g. "contrast mismatch", "wrong modality")
- Proposed Fix: blank for correct rows; your corrected standard name for wrong ones

Medical rules to apply:
- "without contrast" / "plain" in raw_name → standard_name MUST NOT contain "& Contrast" — WRONG if it does. Proposed fix: use the Plain-only name.
- CT in raw_name → standard_name must be CT, not MRI — completely different modality. WRONG if MRI.
- MRI in raw_name → standard_name must be MRI, not CT. WRONG if CT.
- USG/Ultrasound → standard_name must be USG-based. WRONG if X-Ray or CT.
- Body part must match: chest ≠ abdomen, pelvis ≠ spine.
- "Guided" procedures must retain the modality (USG guided, CT guided).

Known errors to catch (examples):
- "CT Scan Chest Without Contrast" → "CT Scan Chest Plain & Contrast" ❌ contrast mismatch → Proposed: "CT Scan Chest Single Plain"
- "C.T scan PELVIS without contrast" → "MRI Pelvis Plain" ❌ modality mismatch → Proposed: "CT Pelvis Plain"
- "CT Temporal Bone Without Contrast" → "CT Temporal Bone Plain & Contrast" ❌ contrast mismatch → Proposed: "CT Temporal Bone Plain"

After your row-by-row analysis:
1. State: "Found X errors, Y look correct."
2. Ask: "Shall I apply all X corrections?" — STOP HERE. No action block.
3. Only after user approval → emit the action block covering all corrections.

### Conversational next-step guidance
After EVERY response — whether you just answered a medical question, made a mapping change, or revalidated mappings — end with a short, natural follow-up suggestion. Keep it to 1–2 sentences. Make it feel like a real conversation, not a bullet list. Examples of tone:
- After a mapping change: "Want me to check the rest of the unmatched rows and suggest fixes for those too?"
- After answering a medical question: "Would you like me to cross-check if any of the tests in this catalogue relate to that condition?"
- After revalidation: "I found X issues and fixed them — want me to do the same for the Radiology tests specifically?"
- After the user just uploads/opens a job: "I can see X unmatched rows — want me to take a first pass at resolving those?"
- If the user seems done: "Looks like you're in good shape — ready to export the final catalogue?"

The suggestion must be DIRECTLY relevant to what was just discussed or what the mapping data shows. Never repeat the same suggestion twice in a row. Keep the tone helpful and brief — like a knowledgeable colleague checking in, not a bot reciting options.

### Current mapping snapshot
"""

    def _strip_price(name: str) -> str:
        return re.sub(r'[\s\-–\t]+R[Ss]\.?\s*[\d,]+\s*$', '', name).strip()

    snapshot = ""

    # Prefer client-supplied mappings (always fresh, works after server restart)
    rows = client_mappings or (jobs[job_id].get("mappings", []) if job_id and job_id in jobs else [])

    if rows:
        total       = len(rows)
        matched_n   = sum(1 for r in rows if r.get("status") in ("matched", "confirmed"))
        unmatched_n = sum(1 for r in rows if r.get("status") == "unmatched")

        def _clean(r: dict) -> dict:
            return {
                "raw_name":      _strip_price(r.get("raw_name", "")),
                "standard_name": r.get("standard_name", ""),
                "match_type":    r.get("match_type", ""),
                "confidence":    round(float(r.get("confidence", 0)), 2),
                "lab_type":      r.get("lab_type", ""),
            }

        snapshot = f"""
Stats: {total} total | {matched_n} matched | {unmatched_n} unmatched

ALL rows (use exact raw_name values to identify rows in action blocks):
{json.dumps([_clean(r) for r in rows], indent=2)}
"""

    return claude_md + persona + snapshot


def _extract_action(text: str) -> dict | None:
    m = re.search(r"```action\s*([\s\S]*?)\s*```", text)
    if m:
        try:
            return json.loads(m.group(1))
        except json.JSONDecodeError:
            pass
    return None


def _strip_action(text: str) -> str:
    return re.sub(r"```action[\s\S]*?```", "", text).strip()


def _rows_to_flat(job: dict) -> list[dict]:
    """Return a single flat list used by the frontend table."""
    rows: list[dict] = []
    for r in job.get("matched", []):
        cat = r.get("catalogue_name", "") or ""
        rows.append({
            "id":             r["id"],
            "raw_name":       r["provider_name"],
            "standard_name":  "" if cat == "nan" else cat,
            "confidence":     r["confidence"],
            "match_type":     r["match_type"],
            "status":         "matched",
            "highlight":      None,
            "price":          r.get("price", ""),
            "lab_type":       r.get("lab_type", ""),
            "serial_no":      r.get("serial_no", ""),
        })
    for r in job.get("unmatched", []):
        cat = r.get("catalogue_name", "") or ""
        rows.append({
            "id":             r["id"],
            "raw_name":       r["provider_name"],
            "standard_name":  "" if cat == "nan" else cat,
            "confidence":     0.0,
            "match_type":     "UNMATCHED",
            "status":         "unmatched",
            "highlight":      None,
            "price":          r.get("price", ""),
            "lab_type":       r.get("lab_type", ""),
            "serial_no":      r.get("serial_no", ""),
        })
    for r in job.get("skipped", []):
        rows.append({
            "id":             r["id"],
            "raw_name":       r["provider_name"],
            "standard_name":  "",
            "confidence":     0.0,
            "match_type":     "SKIPPED",
            "status":         "skipped",
            "highlight":      None,
            "price":          r.get("price", ""),
            "lab_type":       r.get("lab_type", ""),
            "serial_no":      r.get("serial_no", ""),
            "skip_reason":    "Multi-test combination — excluded from output",
        })
    # Sort by source serial_no so order matches the original file
    rows.sort(key=lambda x: int(x["serial_no"]) if str(x.get("serial_no","")).isdigit() else 999999)
    return rows


# ── Claude File Parsing ────────────────────────────────────────────────────────

_PARSE_SYSTEM = """You are a medical lab test catalogue parser following the process-catalogue skill parsing guide.

Your ONLY job: extract raw provider test names from the content provided.

RULES (from parsing-guide.md):
- Return ONLY JSON: {"test_names": [{"name": "...", "price": "...", "lab_type": "Pathology|Radiology|"}]}
- Strip out: prices (Rs/RS + number), codes, units, page numbers, totals
- Include tests from ALL sections regardless of section heading — "OTHER FACILITIES", "SPECIAL TESTS", "RADIOLOGY", "PACKAGES", etc. are all valid sections containing real tests
- Skip rows that are ONLY section headings with no test name (e.g. rows that say only "MRI CHARGES", "TEST NAME:-", "HAEMATOLOGY", "BIO CHEMISTRY", "SEROLOGY" with no actual test)
- Skip column headers (e.g. "Sr No", "Test Name", "Rate", "Amount", "Investigations", "TAT")
- Preserve original casing and spelling — normalization happens in match.py
- Remove duplicates
- Each test name as a separate entry
- Include tests with non-numeric prices like "As per Packages", "Rs 300 per Film", "Rs 300 - Rs 500"
- ORDERING: When the input shows "Left column group" before "Right column group", extract ALL left-column tests first (in their row order), then ALL right-column tests. Do NOT interleave them. This preserves the source file's S.No ordering.
"""


def _pre_extract_excel(file_path: str) -> tuple[list[dict] | None, str]:
    """
    Parsing guide — Excel (multi-sheet aware):

    Pass 1 — deterministic extraction for known BFHL template column headers.
    Scans every sheet for columns whose header matches known provider-name
    patterns.  Also captures price and lab_type (Pathology/Radiology) per row.
    If names are found this way, returns them directly so Claude doesn't need to guess.

    Pass 2 — if Pass 1 finds nothing, falls back to dumping sheet content as
    text so Claude can identify the right column.

    Returns:
        (items, content_text)
        items        — list of {name, price, lab_type} if deterministically found, else None
        content_text — human-readable dump of all sheets (always returned for audit / Claude fallback)
    """
    import pandas as pd
    import math as _math

    # Sheets to skip — admin/meta sheets that never contain test names
    _SKIP_SHEETS = {"provider details", "centre details", "logo", "center details", "instructions", "readme"}

    xl    = pd.ExcelFile(file_path)
    parts = []
    all_items: list[dict] = []
    seen: set[str] = set()

    for sheet in xl.sheet_names:
        if sheet.strip().lower() in _SKIP_SHEETS:
            continue

        df = xl.parse(sheet)
        if df.empty or len(df.columns) < 1:
            continue

        parts.append(f"=== Sheet: {sheet} ({len(df)} rows x {len(df.columns)} cols) ===")
        parts.append("Columns: " + " | ".join(str(c) for c in df.columns))
        for i, (_, row) in enumerate(df.head(3).iterrows()):
            parts.append(f"  Row {i+1}: " + " | ".join(str(v) for v in row.values))
        parts.append("All rows:")
        for _, row in df.iterrows():
            parts.append("  " + " | ".join(str(v) for v in row.values))

        # ── Universal column detection (no hardcoded names) ───────────────────
        col_headers = [str(c) for c in df.columns]
        data_rows   = [[row_data[c] for c in df.columns] for _, row_data in df.iterrows()]
        name_ci, price_ci = _find_name_price_cols(col_headers, data_rows)

        matched_col = df.columns[name_ci]  if name_ci  is not None else None
        price_col   = df.columns[price_ci] if price_ci is not None else None

        # Detect department column
        dept_col = None
        for c in df.columns:
            if str(c).strip().lower() in _DEPT_COLS_SET:
                dept_col = c
                break

        if matched_col is not None:
            # Detect serial-number column to skip section-header rows
            _SR_PAT = re.compile(
                r'^(s\.?\s*[il][\./]?\s*no\.?|sr\.?\s*no\.?|s\.?\s*no\.?|no\.?|#|sn|sl\.?\s*no\.?|serial)$',
                re.IGNORECASE,
            )
            sr_col = None
            for c in df.columns:
                if _SR_PAT.match(str(c).strip()):
                    sr_col = c
                    break
            if sr_col is None:
                for c in df.columns:
                    col_vals = df[c].dropna()
                    num_cnt = sum(1 for v in col_vals if str(v).strip().replace(".", "", 1).isdigit())
                    if num_cnt >= len(col_vals) * 0.5 and len(col_vals) > 0:
                        sr_col = c
                        break

            for _, row_data in df.iterrows():
                raw_val = row_data.get(matched_col)
                if raw_val is None or (isinstance(raw_val, float) and _math.isnan(raw_val)):
                    continue
                name = str(raw_val).strip()
                if not name or name.lower() in seen or name.lower() in ("nan",):
                    continue
                if sr_col is not None:
                    sr_val = row_data.get(sr_col)
                    sr_empty = sr_val is None or (isinstance(sr_val, float) and _math.isnan(sr_val)) or str(sr_val).strip() in ("", "nan")
                    if sr_empty:
                        continue
                # Price
                price_val = ""
                if price_col is not None:
                    pv = row_data.get(price_col)
                    if pv is not None and not (isinstance(pv, float) and _math.isnan(pv)):
                        pv_s = str(pv).strip()
                        if pv_s and pv_s != "nan":
                            price_val = pv_s
                # Lab type
                lab_type = _infer_lab_type(name)
                if dept_col is not None:
                    dv = str(row_data.get(dept_col, "") or "").strip()
                    if dv and dv.lower() not in ("nan", "") and _RAD_PAT.search(dv):
                        lab_type = "Radiology"

                seen.add(name.lower())
                all_items.append({"name": name, "price": price_val, "lab_type": lab_type})

    content = "\n".join(parts)
    return (all_items if all_items else None, content)


def _pre_extract_pdf(file_path: str) -> str:
    """
    Parsing guide — PDF:
    Try table extraction first (most catalogues are table-formatted) so
    columns stay aligned; fall back to plain text per page.
    Claude then identifies lines that are test names and strips prices,
    codes, section headers, page numbers.
    """
    import pdfplumber
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            page_parts = []

            # ── Try tables first ──────────────────────────────────────────────
            tables = page.extract_tables() or []
            for ti, table in enumerate(tables):
                page_parts.append(f"  [Table {ti}]")
                for row in table:
                    if row:
                        page_parts.append("  " + " | ".join(
                            str(cell).strip() if cell else "" for cell in row
                        ))

            # ── Fall back to plain text if no tables found on this page ───────
            if not tables:
                t = page.extract_text()
                if t:
                    page_parts.append(t)

            if page_parts:
                pages.append(f"--- Page {i} ---\n" + "\n".join(page_parts))

    return "\n".join(pages)


_SNO_HDR_RE = re.compile(
    r'^(s\.?\s*[il][\./]?\s*no\.?|sr\.?\s*no\.?|s\.?\s*no\.?|serial|sno|no\.?|#|sl\.?\s*no\.?)$',
    re.I,
)


def _pre_extract_docx(file_path: str) -> str:
    """
    Parsing guide — DOCX:
    Extract from tables first (most catalogues are table-formatted),
    then paragraphs.

    Two-column table handling (common in printed rate lists):
    If the table has an S.No column on both the left half and right half,
    extract LEFT column group first (all rows), then RIGHT column group.
    This preserves the source file's S.No ordering (1,2,3…60 then 61,62…).
    The S.No value is included in each row so Claude can carry it through.
    """
    import re as _re
    from docx import Document
    doc   = Document(file_path)
    parts = []

    for ti, table in enumerate(doc.tables):
        rows_data = [[c.text.strip() for c in row.cells] for row in table.rows]
        if not rows_data:
            continue
        n_cols = len(rows_data[0])
        parts.append(f"=== Table {ti} ({len(rows_data)} rows x {n_cols} cols) ===")

        # Show header row for structure context
        parts.append(f"  Header: {rows_data[0]}")

        # ── Detect two-column (side-by-side) table ────────────────────────────
        # Pattern: 4-6 cols where col 0 and col half both look like S.No headers
        half = n_cols // 2
        hdr  = [c.lower().strip() for c in rows_data[0]]
        is_double = (
            n_cols >= 4
            and half < len(hdr)
            and _SNO_HDR_RE.match(hdr[0])
            and _SNO_HDR_RE.match(hdr[half])
        )

        if is_double:
            # Left group: cols 0 … half-1
            parts.append(f"-- Left column group (cols 0-{half-1}):")
            parts.append("  " + " | ".join(rows_data[0][:half]))  # header
            for row in rows_data[1:]:
                cells = row[:half]
                if any(c for c in cells):
                    parts.append("  " + " | ".join(cells))
            # Right group: cols half … n_cols-1
            parts.append(f"-- Right column group (cols {half}-{n_cols-1}):")
            parts.append("  " + " | ".join(rows_data[0][half:]))  # header
            for row in rows_data[1:]:
                cells = row[half:]
                if any(c for c in cells):
                    parts.append("  " + " | ".join(cells))
        else:
            parts.append("All rows:")
            for row in rows_data:
                parts.append("  " + " | ".join(row))

    # Paragraphs
    if doc.paragraphs:
        parts.append("=== Paragraphs ===")
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

    return "\n".join(parts)


def _pre_extract_image(file_path: str) -> bytes:
    """Return raw image bytes for Claude vision."""
    return Path(file_path).read_bytes()


def _parse_file_with_claude(file_path: str, job_id: str | None = None) -> list[dict]:
    """
    Use Claude Sonnet to extract test names following the skill parsing guide.
    Each file type is pre-processed in Python per guide instructions,
    then Claude identifies and cleans the test names.
    Times out after 90s to avoid hanging indefinitely.

    Returns list of dicts: [{name, price, lab_type}, ...]
    """
    if not os.environ.get("ANTHROPIC_API_KEY"):
        return []

    path = Path(file_path)
    ext  = path.suffix.lower()

    def _prog(pct: int, step: str):
        if job_id:
            _progress(job_id, pct, step)

    try:
        # ── Step 1: Pre-process per guide ────────────────────────────────────
        if ext in (".xlsx", ".xls"):
            _prog(7, "Reading Excel sheets…")
            det_names, content = _pre_extract_excel(file_path)

            if det_names:
                # Deterministic extraction succeeded — skip Claude for parsing
                _prog(25, f"Extracted {len(det_names)} items from Excel (deterministic)…")
                return det_names

            # Fallback: ask Claude to identify the right column
            prompt  = (
                "This is an Excel provider lab test catalogue — it may have multiple sheets.\n"
                "Your task: extract EVERY test/investigation name from ALL sheets that contain test data.\n"
                "Ignore admin sheets (e.g. Provider Details, Instructions, Logo).\n\n"
                "Rules — independent of column names:\n"
                "1. The TEST NAME column contains medical test/investigation names (text, 3–80 chars). "
                "Find it by content, not by header. The header could say anything.\n"
                "2. The PRICE column contains numeric values (e.g. 450, 1200). "
                "Find it by content — the header could say 'Rate', 'MRP', 'Amount', 'Charges', or anything else. "
                "ALWAYS extract price — do not skip it even if the header is unusual.\n"
                "3. The DEPARTMENT column (if any) says 'Pathology' or 'Radiology'.\n"
                "4. Skip: row numbers, codes, type labels (Routine/Special), specimen types "
                "(Serum/Blood/Urine), TAT values, section headers, totals.\n\n"
                f"{content[:80_000]}\n\n"
                'Return JSON: {"tests": [{"name": "Complete Blood Count", "price": "450", "lab_type": "Pathology"}, ...]}'
            )
            messages = [{"role": "user", "content": prompt}]

        elif ext == ".pdf":
            _prog(7, "Extracting PDF text…")

            # ── Pass 1: universal column detection ───────────────────────────
            _SKIP_VALUES = {
                "routine", "special", "serum", "plasma", "urine", "blood",
                "edta blood", "pus", "fluid", "sputum", "stool", "swab",
                "slides", "smear", "same day", "next day", "yes", "no",
                "type", "specimen", "tat", "price", "code", "sr no", "s no",
            }

            import pdfplumber
            det_items: list[dict] = []
            seen_det: set[str] = set()
            global_col_idx: int | None = None
            global_price_idx: int | None = None
            global_dept_idx: int | None = None

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    for table in (page.extract_tables() or []):
                        if not table:
                            continue

                        # Detect columns using universal scorer on first 5 rows as header candidates
                        header_row_idx = 0
                        header_col_idx = None
                        price_idx: int | None = None
                        dept_idx: int | None = None

                        for ri, hrow in enumerate(table[:5]):
                            if not hrow:
                                continue
                            hdrs = [str(c or "") for c in hrow]
                            data_sample = [r for r in table[ri + 1: ri + 15] if r]
                            nc, pc = _find_name_price_cols(hdrs, data_sample)
                            if nc is not None:
                                header_row_idx = ri
                                header_col_idx = nc
                                price_idx      = pc
                                # Dept col: scan header row for dept keyword
                                for ci2, cell2 in enumerate(hrow):
                                    if str(cell2 or "").strip().lower() in _DEPT_COLS_SET:
                                        dept_idx = ci2
                                        break
                                break

                        # If no header found, reuse last detected columns
                        if header_col_idx is None and global_col_idx is not None:
                            header_col_idx = global_col_idx
                            price_idx      = global_price_idx
                            dept_idx       = global_dept_idx
                            header_row_idx = 0

                        if header_col_idx is not None:
                            global_col_idx   = header_col_idx
                            global_price_idx = price_idx
                            global_dept_idx  = dept_idx
                            for row in table[header_row_idx + 1:]:
                                if not row or header_col_idx >= len(row):
                                    continue
                                val = str(row[header_col_idx] or "").strip()
                                val_low = val.lower()
                                if (val and len(val) >= 3
                                        and val_low not in _SKIP_VALUES
                                        and val_low not in seen_det
                                        and not val.replace(" ", "").isdigit()):
                                    # Price
                                    price_val = ""
                                    if price_idx is not None and price_idx < len(row):
                                        pv = str(row[price_idx] or "").strip()
                                        if pv and pv.lower() not in ("price", "rate", "nan", ""):
                                            price_val = pv
                                    # Lab type
                                    lab_type = _infer_lab_type(val)
                                    if dept_idx is not None and dept_idx < len(row):
                                        dv = str(row[dept_idx] or "").strip()
                                        if dv and _RAD_PAT.search(dv):
                                            lab_type = "Radiology"
                                    seen_det.add(val_low)
                                    det_items.append({"name": val, "price": price_val, "lab_type": lab_type})

            if det_items:
                _prog(25, f"Extracted {len(det_items)} items from PDF table (deterministic)…")
                return det_items

            # ── Pass 2: Claude text / vision fallback ─────────────────────────
            content = _pre_extract_pdf(file_path)

            if content.strip():
                # Text-based PDF — send as structured text
                _prog(14, "Identifying test names with Claude…")
                prompt = (
                    "This is a PDF provider lab test catalogue.\n"
                    "Your task: extract EVERY test/investigation name along with its price.\n\n"
                    "Rules — independent of column names:\n"
                    "1. The TEST NAME column contains medical test names (text, 3–80 chars). "
                    "Find it by content — the header could say anything (Test, Investigation, Description, Parameter, Item, Service, etc.).\n"
                    "2. The PRICE column contains numeric values (e.g. 450, 1200). "
                    "Find it by content — ALWAYS extract it regardless of what the header says. "
                    "The header might be Rate, MRP, Charges, Amount, Fee, Rs, or anything else.\n"
                    "3. If a DEPARTMENT or SECTION column exists, classify each test as Pathology or Radiology.\n"
                    "4. Skip: row numbers, codes, type labels (Routine/Special), specimen types "
                    "(Serum/Blood/Urine/Fluid), TAT values, section headers, page numbers.\n\n"
                    f"{content[:80_000]}\n\n"
                    'Return JSON: {"tests": [{"name": "Complete Blood Count", "price": "450", "lab_type": "Pathology"}, ...]}'
                )
                messages = [{"role": "user", "content": prompt}]
            else:
                # Scanned / image-based PDF — process one page at a time (smaller requests)
                _prog(10, "Scanned PDF detected — processing pages with Claude vision…")
                import base64
                try:
                    import fitz  # PyMuPDF
                except ImportError:
                    print("WARNING: PyMuPDF not installed — pip install pymupdf", file=sys.stderr)
                    return []

                doc = fitz.open(file_path)
                all_page_items: list[dict] = []
                seen_vision: set[str] = set()
                total_pages = min(len(doc), 40)
                BATCH_SIZE = 4   # pages per API call

                # Render all pages to JPEG first
                page_images: list[tuple[int, str]] = []  # (page_num, b64)
                for page_num, page in enumerate(doc, 1):
                    if page_num > total_pages:
                        break
                    mat = fitz.Matrix(1.5, 1.5)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img_bytes = pix.tobytes("jpeg", jpg_quality=75)
                    page_images.append((page_num, base64.standard_b64encode(img_bytes).decode()))
                doc.close()

                # Process in batches
                batches = [page_images[i:i+BATCH_SIZE] for i in range(0, len(page_images), BATCH_SIZE)]
                for batch_idx, batch in enumerate(batches):
                    page_nums = [p for p, _ in batch]
                    _prog(10 + int((batch_idx + 1) / len(batches) * 12),
                          f"Vision batch {batch_idx+1}/{len(batches)} (pages {page_nums[0]}–{page_nums[-1]})…")

                    content_parts: list[dict] = []
                    for page_num, img_b64 in batch:
                        content_parts.append({"type": "text", "text": f"--- Page {page_num} ---"})
                        content_parts.append({"type": "image", "source": {
                            "type": "base64", "media_type": "image/jpeg", "data": img_b64}})

                    batch_prompt = (
                        "These are consecutive pages from a scanned provider lab test catalogue. "
                        "Extract EVERY test/investigation name from ALL pages — do not skip any row. "
                        "For PRICE: find the numeric column regardless of its header name "
                        "(could be Rate, MRP, Charges, Amount, Fee, Rs, or anything else). "
                        "ALWAYS capture the price for each test. "
                        "Classify each test as Pathology or Radiology. "
                        "Ignore logos, headers, footers, page numbers, and section dividers. "
                        'Return JSON: {"tests": [{"name": "CT Head Plain", "price": "6000", "lab_type": "Radiology"}, ...]}'
                    )
                    content_parts.append({"type": "text", "text": batch_prompt})

                    try:
                        batch_resp = claude.messages.create(
                            model="claude-opus-4-6",
                            max_tokens=4096,
                            messages=[{"role": "user", "content": content_parts}],
                            timeout=120.0,
                        )
                        raw_batch = batch_resp.content[0].text
                        pm = re.search(r'\{[\s\S]*\}', raw_batch)
                        if pm:
                            pd_ = json.loads(pm.group(0))
                            for item in pd_.get("tests", []):
                                n = str(item.get("name", "")).strip()
                                if n and n.lower() not in seen_vision:
                                    seen_vision.add(n.lower())
                                    lab_type = str(item.get("lab_type", "")).strip()
                                    if lab_type not in ("Pathology", "Radiology"):
                                        lab_type = _infer_lab_type(n)
                                    price = str(item.get("price", "") or "").strip()
                                    if price.lower() in ("nan","none","null","—","-","n/a",""):
                                        price = ""
                                    all_page_items.append({"name": n, "price": price, "lab_type": lab_type})
                        print(f"[INFO] Batch {batch_idx+1}: extracted {len(pd_.get('tests',[]) if pm else [])} tests", flush=True)
                    except Exception as _batch_exc:
                        _bmsg = str(_batch_exc).lower()
                        if "credit balance is too low" in _bmsg or "your credit balance" in _bmsg:
                            print(f"[WARN] Credit limit hit on batch {batch_idx+1} — stopping vision pass", flush=True)
                            break
                        print(f"[WARN] Batch {batch_idx+1} vision failed: {_batch_exc}", flush=True)
                        continue

                if all_page_items:
                    _prog(25, f"Extracted {len(all_page_items)} items from scanned PDF…")
                    return all_page_items
                return []

        elif ext in (".docx", ".doc"):
            _prog(7, "Reading Word document tables…")

            # ── Pass 1: deterministic column match ────────────────────────────
            _PARAM_COLS_W = [
                "parameter descriptions", "parameter description",
                "test name", "test names", "investigation", "investigations",
                "item name", "item", "description", "service", "procedure",
                "test", "particular", "name",
            ]
            _SKIP_VALUES_W = {
                "routine", "special", "serum", "plasma", "urine", "blood",
                "edta blood", "pus", "fluid", "sputum", "stool", "swab",
                "slides", "smear", "same day", "next day", "yes", "no",
                "type", "specimen", "tat", "price", "code", "sr no", "s no",
            }
            try:
                from docx import Document as _DocxDoc
                _doc = _DocxDoc(file_path)
                _det_items: list[dict] = []
                _seen_w: set[str] = set()

                for table in _doc.tables:
                    if not table.rows:
                        continue
                    # Find header row — detect ALL name cols (supports two-column layouts)
                    _hcols: list[int] = []   # all name column indices
                    _pcols: dict[int, int] = {}  # name_col -> price_col
                    _hrow  = 0
                    _dcol  = None
                    _hdr_cells_low: list[str] = []
                    _NAME_KWS_W = ["investigation", "name", "test", "item", "description",
                                   "parameter", "service", "procedure", "particular"]

                    # Pass A: keyword-based detection (handles labelled headers + embedded headers)
                    for ri, row in enumerate(table.rows[:5]):
                        cells_text = [cell.text.strip() for cell in row.cells]
                        cells_low  = [t.lower() for t in cells_text]
                        for ci, cl in enumerate(cells_low):
                            is_name_col = any(k in cl for k in _NAME_KWS_W)
                            if is_name_col and ci not in _hcols:
                                _hcols.append(ci)
                                _hrow = ri
                            if _dcol is None and cl in _DEPT_COLS_SET:
                                _dcol = ci
                        if _hcols:
                            _hdr_cells_low = cells_low
                            # Find nearest price col for each name col (prefer rightward)
                            for nc in _hcols:
                                best_pc, best_dist = None, 99
                                for pc, cl in enumerate(cells_low):
                                    if not _is_price_col(cl):
                                        continue
                                    dist = pc - nc
                                    if 0 < dist <= 4 and dist < best_dist:
                                        best_pc, best_dist = pc, dist
                                if best_pc is None:
                                    for pc, cl in enumerate(cells_low):
                                        if _is_price_col(cl) and abs(pc - nc) <= 4:
                                            best_pc = pc
                                            break
                                if best_pc is not None:
                                    _pcols[nc] = best_pc
                            break

                    # Deduplicate: keep only cols whose header actually contains a name keyword
                    _hcols = [nc for nc in _hcols
                              if nc < len(_hdr_cells_low)
                              and any(k in _hdr_cells_low[nc] for k in _NAME_KWS_W)]

                    # Pass B: content-based fallback when keyword detection found nothing
                    if not _hcols and len(table.rows) >= 3:
                        all_cell_rows = [[cell.text.strip() for cell in row.cells]
                                         for row in table.rows]
                        hdrs = all_cell_rows[0]
                        fb_nc, fb_pc = _find_name_price_cols(hdrs, all_cell_rows[1:])
                        if fb_nc is not None:
                            _hcols = [fb_nc]
                            _hrow  = 0
                            _hdr_cells_low = hdrs
                            if fb_pc is not None:
                                _pcols[fb_nc] = fb_pc

                    if _hcols:
                        # Find S.No column for each name column (first S.No-looking col before it)
                        _sno_for_nc: dict[int, int] = {}
                        for _nc in _hcols:
                            for _sci in range(_nc):
                                _sh = _hdr_cells_low[_sci] if _sci < len(_hdr_cells_low) else ""
                                if _SNO_HDR_RE.match(_sh):
                                    _sno_for_nc[_nc] = _sci
                                    break

                        # KEY FIX: iterate name columns FIRST, then rows — so left column
                        # items all come before right column items (preserves S.No ordering).
                        for nc in _hcols:
                            sno_col = _sno_for_nc.get(nc)
                            for row in table.rows[_hrow:]:
                                ncells = len(row.cells)
                                if nc >= ncells:
                                    continue
                                val = row.cells[nc].text.strip()
                                # Strip embedded header prefix (e.g. "Name of Investigation Actual Name"
                                # or "Name of Investigation\nActual Name")
                                if val.lower().startswith("name of investigation"):
                                    val = val[len("name of investigation"):].strip().lstrip("\n").strip()
                                # Extract numeric price from cells that embed "Charges\n500"
                                price_val = ""
                                pc = _pcols.get(nc)
                                if pc is not None and pc < ncells:
                                    pv = row.cells[pc].text.strip()
                                    # Handle "Charges\n500" pattern
                                    if "\n" in pv:
                                        pv = pv.split("\n")[-1].strip()
                                    if pv and pv.lower() not in ("price", "rate", "charges", "nan", ""):
                                        price_val = pv
                                # Extract S.No value for use as serial_no
                                sno_val = ""
                                if sno_col is not None and sno_col < ncells:
                                    sno_val = row.cells[sno_col].text.strip()
                                val_low = val.lower()
                                if (val and len(val) >= 3
                                        and val_low not in _SKIP_VALUES_W
                                        and val_low not in _seen_w
                                        and not val.replace(" ", "").isdigit()):
                                    lab_type = _infer_lab_type(val)
                                    if _dcol is not None and _dcol < ncells:
                                        dv = row.cells[_dcol].text.strip()
                                        if dv and _RAD_PAT.search(dv):
                                            lab_type = "Radiology"
                                    _seen_w.add(val_low)
                                    item = {"name": val, "price": price_val, "lab_type": lab_type}
                                    if sno_val and sno_val.isdigit():
                                        item["serial_no"] = int(sno_val)
                                    _det_items.append(item)

                if _det_items:
                    _prog(25, f"Extracted {len(_det_items)} items from Word table (deterministic)…")
                    return _det_items
            except Exception:
                pass  # fall through to Claude

            # ── Pass 2: Claude fallback ───────────────────────────────────────
            content = _pre_extract_docx(file_path)
            _prog(14, "Identifying test names with Claude…")
            prompt  = (
                "This is a Word document provider lab test catalogue.\n"
                "Your task: extract EVERY test/investigation name along with its price.\n\n"
                "Rules — independent of column names:\n"
                "1. The TEST NAME column contains medical test names (text, 3–80 chars). "
                "Find it by content — the header could say anything (Test, Investigation, Description, Parameter, Item, Service, etc.).\n"
                "2. The PRICE column contains numeric values. ALWAYS extract it regardless of the header name "
                "(Rate, MRP, Charges, Amount, Fee, Rs, Cost — any of these).\n"
                "3. If a DEPARTMENT or SECTION column exists, classify each test as Pathology or Radiology.\n"
                "4. Skip: row numbers, codes, type labels (Routine/Special), specimen types "
                "(Serum/Blood/Urine/Fluid), TAT values, section headers, footnotes.\n\n"
                f"{content[:80_000]}\n\n"
                'Return JSON: {"tests": [{"name": "Complete Blood Count", "price": "450", "lab_type": "Pathology"}, ...]}'
            )
            messages = [{"role": "user", "content": prompt}]

        elif ext in (".jpg", ".jpeg", ".png", ".webp"):
            _prog(7, "Reading image with Claude vision…")
            import base64
            img_bytes = _pre_extract_image(file_path)
            img_b64   = base64.standard_b64encode(img_bytes).decode()
            media_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                         ".png": "image/png",  ".webp": "image/webp"}
            messages = [{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64",
                    "media_type": media_map.get(ext, "image/jpeg"), "data": img_b64}},
                {"type": "text", "text": (
                    "This is an image of a provider lab test catalogue.\n"
                    "Your task: extract EVERY test/investigation name along with its price.\n\n"
                    "Rules — independent of column names:\n"
                    "1. Find the TEST NAME column by its content (medical test names, 3–80 chars). "
                    "The column header could say anything.\n"
                    "2. Find the PRICE column by its content (numeric values like 450, 1200). "
                    "ALWAYS extract it — the header might be Rate, MRP, Charges, Amount, Fee, or anything else.\n"
                    "3. Classify each test as Pathology or Radiology if a department column is visible.\n"
                    "4. Skip: row/serial numbers, test codes, type labels (Routine/Special), "
                    "specimen types (Serum/Blood/Urine/Fluid/Stool/Swab), TAT values.\n"
                    'Return JSON: {"tests": [{"name": "Complete Blood Count", "price": "450", "lab_type": "Pathology"}, ...]}'
                )},
            ]}]

        elif ext == ".csv":
            _prog(7, "Reading CSV…")
            content  = path.read_text(encoding="utf-8", errors="ignore")
            _prog(14, "Identifying test names with Claude…")
            prompt   = (
                "This is a CSV provider lab test catalogue.\n"
                "Your task: extract EVERY test/investigation name along with its price.\n\n"
                "Rules — independent of column names:\n"
                "1. Find the TEST NAME column by content (medical test names, 3–80 chars). Header could say anything.\n"
                "2. Find the PRICE column by content (numeric values). ALWAYS extract it regardless of header name.\n"
                "3. Classify as Pathology or Radiology if a department column exists.\n"
                "4. Skip: row numbers, codes, type labels, specimen types, TAT, section headers.\n\n"
                f"{content[:80_000]}\n\n"
                'Return JSON: {"tests": [{"name": "Complete Blood Count", "price": "450", "lab_type": "Pathology"}, ...]}'
            )
            messages = [{"role": "user", "content": prompt}]

        else:
            return []

        # ── Step 2: Claude extracts test names (90s timeout) ─────────────────
        _prog(18, "Parsing with Claude…")
        resp = claude.messages.create(
            model="claude-opus-4-6",
            max_tokens=4096,
            system=_PARSE_SYSTEM,
            messages=messages,
            timeout=90.0,
        )
        raw  = resp.content[0].text
        m    = re.search(r'\{[\s\S]*\}', raw)
        if not m:
            return []

        data = json.loads(m.group(0))

        # Support both new {"tests": [{name, price, lab_type}]} and legacy {"test_names": [...]}
        raw_tests = data.get("tests", [])
        if not raw_tests:
            # Legacy fallback — flat name list, no price/lab_type from Claude
            raw_tests = [{"name": str(n).strip(), "price": "", "lab_type": ""}
                         for n in data.get("test_names", []) if str(n).strip()]

        # ── Step 3: Quality checks (parsing-guide.md) ────────────────────────
        seen: set[str] = set()
        unique: list[dict] = []
        for item in raw_tests:
            n = str(item.get("name", "")).strip()
            if not n or n.lower() in seen:
                continue
            seen.add(n.lower())
            lab_type = str(item.get("lab_type", "")).strip()
            if lab_type not in ("Pathology", "Radiology"):
                lab_type = _infer_lab_type(n)
            price = str(item.get("price", "") or "").strip()
            if price.lower() in ("nan", "none", "null", "—", "-", "n/a"):
                price = ""
            unique.append({"name": n, "price": price, "lab_type": lab_type})

        return unique

    except Exception as _exc:
        import traceback; traceback.print_exc()
        print(f"[DEBUG] _parse_file_with_claude EXCEPTION: {type(_exc).__name__}: {_exc}", flush=True)
        # Surface billing / auth errors directly — match Anthropic's exact phrases only
        # (broad matches like "billing"/"insufficient" falsely catch pdfplumber/PDF errors)
        _msg = str(_exc).lower()
        _is_billing = (
            "credit balance is too low" in _msg
            or "your credit balance" in _msg
            or "insufficient_quota" in _msg
            or ("quota" in _msg and "anthropic" in _msg)
        )
        if _is_billing:
            # Non-fatal: log warning and fall back to deterministic parser
            # (Claude API workspace spending limit may be low even if org balance is positive)
            print(f"[WARN] Anthropic API credit/quota limit hit — falling back to deterministic parser", flush=True)
            if job_id:
                _progress(job_id, 8, "Claude unavailable (quota) — using deterministic parser…")
            return []
        if "api key" in _msg or "authentication" in _msg or "unauthorized" in _msg or "x-api-key" in _msg:
            print(f"[WARN] Anthropic API key invalid — falling back to deterministic parser", flush=True)
            if job_id:
                _progress(job_id, 8, "Claude unavailable (auth) — using deterministic parser…")
            return []
        # All other errors: log and fall through so the fallback parser can handle it
        print(f"[DEBUG] _parse_file_with_claude non-billing error — falling back to processor.py", flush=True)
        return []


# ── Semantic Recovery ──────────────────────────────────────────────────────────

def _load_skill_context() -> str:
    """Load accuracy-loop-guide.md for semantic recovery system prompt."""
    refs_dir = PROJECT_ROOT / ".claude" / "skills" / "process-catalogue" / "references"
    parts = []
    loop_guide = refs_dir / "accuracy-loop-guide.md"
    if loop_guide.exists():
        parts.append(loop_guide.read_text(encoding="utf-8"))
    parts.append(
        "\nYou are performing the semantic recovery pass (from accuracy-loop-guide.md).\n"
        "Match each unmatched provider test name to the best catalogue name from the candidates provided.\n"
        "Return ONLY valid JSON: "
        '{"matches": [{"id": "...", "catalogue_name": "exact name or null", "confidence": 0.75, "skipped": false}]}\n'
        "Only include rows you matched. Use null for catalogue_name if no confident match (≥65%) exists.\n"
        "Set skipped=true for combination tests (A & B / A AND B / A WITH B with two distinct tests)."
    )
    return "\n\n---\n\n".join(parts)


def _master_db_path() -> Path:
    return PROJECT_ROOT / "refrences" / "Master.csv.db"


def _load_catalogue_names() -> list[str]:
    db_path = _master_db_path()
    if not db_path.exists():
        return []
    try:
        conn = sqlite3.connect(str(db_path))
        rows = conn.execute("SELECT DISTINCT catalogue_name FROM master").fetchall()
        conn.close()
        return [r[0] for r in rows if r[0]]
    except Exception:
        return []


def _get_candidates(name: str, catalogue_names: list[str], n: int = 20) -> list[str]:
    if not catalogue_names:
        return []
    results = fuzz_process.extract(name, catalogue_names, scorer=fuzz.token_sort_ratio, limit=n)
    return [r[0] for r in results if r[1] >= 35]


def _semantic_recovery(unmatched: list[dict]) -> tuple[list[dict], list[dict]]:
    """
    Use Claude Sonnet (following skill rules + CLAUDE.md) to semantically
    recover unmatched rows. GPT-4o is NOT used here.
    Returns (newly_matched, still_unmatched).
    """
    if not unmatched:
        return [], []

    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("⚠️  ANTHROPIC_API_KEY not set — skipping semantic recovery")
        return [], unmatched

    catalogue_names   = _load_catalogue_names()
    system_prompt     = _load_skill_context()
    recovered_matched: list[dict] = []
    still_unmatched:   list[dict] = []

    BATCH = 25
    for i in range(0, len(unmatched), BATCH):
        batch = unmatched[i : i + BATCH]

        rows_payload = []
        for row in batch:
            candidates = _get_candidates(row["provider_name"], catalogue_names)
            rows_payload.append({
                "id":         row["id"],
                "name":       row["provider_name"],
                "candidates": candidates,
            })

        prompt = (
            "Apply the semantic recovery pass (accuracy-loop-guide.md Pass 2) to these unmatched rows.\n"
            "Each row has fuzzy candidates pre-computed. Use your medical knowledge to pick the best match.\n\n"
            f"{json.dumps(rows_payload, indent=2)}\n\n"
            'Return JSON: {"matches": [{"id":"...","catalogue_name":"exact name or null","confidence":0.0,"skipped":false}]}'
        )

        try:
            resp = claude.messages.create(
                model="claude-opus-4-6",
                max_tokens=3000,
                system=system_prompt,
                messages=[{"role": "user", "content": prompt}],
            )
            raw  = resp.content[0].text
            m    = re.search(r'\{[\s\S]*\}', raw)
            data = json.loads(m.group(0)) if m else {}
            matches = {m_["id"]: m_ for m_ in data.get("matches", [])}
        except Exception:
            import traceback; traceback.print_exc()
            matches = {}

        for row in batch:
            m = matches.get(row["id"])
            if m and m.get("skipped"):
                row["match_type"] = "SKIPPED"
                still_unmatched.append(row)
            elif m and m.get("catalogue_name") and float(m.get("confidence", 0)) >= 0.65:
                row["catalogue_name"] = m["catalogue_name"]
                row["confidence"]     = float(m["confidence"])
                row["match_type"]     = "fuzzy-semantic"
                recovered_matched.append(row)
            else:
                still_unmatched.append(row)

    return recovered_matched, still_unmatched


# ── Upload & Process ───────────────────────────────────────────────────────────
@app.post("/api/upload")
async def upload(file: UploadFile = File(...), background_tasks: BackgroundTasks = BackgroundTasks()):
    import tempfile
    job_id  = str(uuid.uuid4())
    tmp_dir = Path(tempfile.gettempdir()) / "console_jobs" / job_id
    tmp_dir.mkdir(parents=True, exist_ok=True)

    file_path = tmp_dir / (file.filename or "upload")
    file_path.write_bytes(await file.read())

    jobs[job_id] = {
        "status":       "processing",
        "progress":     0,
        "current_step": "Queued…",
        "filename":     file.filename or "upload",
        "file_path":    str(file_path),
        "matched":      [],
        "unmatched":    [],
        "skipped":      [],
        "mappings":     [],
        "stats":        {},
        "output_path":  None,
        "error":        None,
    }
    sessions[job_id] = []

    background_tasks.add_task(_process_job, job_id, str(file_path))
    return {"job_id": job_id, "filename": file.filename}


def _progress(job_id: str, pct: int, step: str):
    if job_id in jobs:
        jobs[job_id]["progress"] = pct
        jobs[job_id]["current_step"] = step


def _process_job(job_id: str, file_path: str):
    try:
        _progress(job_id, 5, "Extracting text from file…")
        print(f"[DEBUG] _parse_file_with_claude starting for: {file_path}", flush=True)
        parsed_items = _parse_file_with_claude(file_path, job_id)
        print(f"[DEBUG] _parse_file_with_claude returned {len(parsed_items)} items", flush=True)
        if not parsed_items:
            # Fallback to processor.py if Claude parsing fails/times out
            print(f"[DEBUG] FALLBACK triggered — using parse_file", flush=True)
            _progress(job_id, 8, "Falling back to local parser…")
            plain_names = parse_file(file_path, jobs, job_id)
            print(f"[DEBUG] parse_file returned {len(plain_names)} names", flush=True)
            parsed_items = [{"name": n, "price": "", "lab_type": _infer_lab_type(n)} for n in plain_names]
        if not parsed_items:
            raise ValueError("No test names could be extracted.")

        # Stamp sequential serial_no preserving source file order
        for idx, item in enumerate(parsed_items):
            item.setdefault("serial_no", idx + 1)

        names    = [r["name"] for r in parsed_items]
        meta_map = {r["name"]: r for r in parsed_items}  # name → {price, lab_type, serial_no}

        _progress(job_id, 30, f"Extracted {len(names)} names — running match.py…")
        results = run_match_script(names, PROJECT_ROOT, job_id, jobs)

        _progress(job_id, 80, "Categorising results…")
        matched, unmatched, skipped = [], [], []
        for row in results:
            mt         = str(row.get("Match Type", "UNMATCHED"))
            prov_name  = str(row.get("Provider Test Name", ""))
            meta       = meta_map.get(prov_name, {})
            item = {
                "id":             str(uuid.uuid4()),
                "provider_name":  prov_name,
                "catalogue_name": "" if (v := row.get("Catalogue Test Name")) is None or (isinstance(v, float) and __import__('math').isnan(v)) else str(v),
                "match_type":     mt,
                "confidence":     float(row.get("Confidence Score") or 0),
                "price":          meta.get("price", ""),
                "lab_type":       meta.get("lab_type", "") or _infer_lab_type(prov_name),
                "serial_no":      meta.get("serial_no", ""),
            }
            if mt == "SKIPPED":
                skipped.append(item)
            elif mt == "UNMATCHED":
                unmatched.append(item)
            else:
                matched.append(item)

        # ── Apply past user corrections ───────────────────────────────────────
        corr_path = PROJECT_ROOT / "learning" / "corrections.json"
        if corr_path.exists():
            try:
                corrections_data = json.loads(corr_path.read_text(encoding="utf-8"))
                # Build lookup: normalized provider name → {standard_name, lab_type, price}
                corr_lookup: dict[str, dict] = {}
                for c in corrections_data:
                    pn = str(c.get("provider_name", "")).strip().lower()
                    if pn and c.get("new_catalogue_name"):
                        corr_lookup[pn] = {
                            "standard_name": c["new_catalogue_name"],
                            "lab_type":      c.get("lab_type", ""),
                            "price":         c.get("price", ""),
                        }
                # Apply to all items (matched + unmatched)
                for item in matched + unmatched:
                    key = item.get("provider_name", "").strip().lower()
                    if key in corr_lookup:
                        override = corr_lookup[key]
                        item["standard_name"]  = override["standard_name"]
                        item["catalogue_name"] = override["standard_name"]
                        item["match_type"]     = "correction"
                        item["confidence"]     = 1.0
                        if override.get("lab_type"):
                            item["lab_type"] = override["lab_type"]
                        if override.get("price"):
                            item["price"] = override["price"]
            except Exception as _ce:
                print(f"[WARN] Could not apply corrections: {_ce}", flush=True)

        # ── Semantic recovery pass ────────────────────────────────────────────
        if unmatched:
            _progress(job_id, 85, f"Semantic recovery for {len(unmatched)} unmatched rows…")
            recovered, still_unmatched = _semantic_recovery(unmatched)
            # Move SKIPPED-flagged rows out of still_unmatched
            new_skipped   = [r for r in still_unmatched if r["match_type"] == "SKIPPED"]
            still_unmatched = [r for r in still_unmatched if r["match_type"] != "SKIPPED"]
            matched   += recovered
            skipped   += new_skipped
            unmatched  = still_unmatched
            _progress(job_id, 95, f"Recovered {len(recovered)} rows semantically")

        stats = {
            "total":     len(names),
            "matched":   len(matched),
            "unmatched": len(unmatched),
            "skipped":   len(skipped),
            "exact":     sum(1 for r in results if r.get("Match Type") == "exact"),
            "fuzzy":     sum(1 for r in results if "fuzzy" in str(r.get("Match Type", ""))),
            "semantic":  sum(1 for r in matched if r.get("match_type") == "fuzzy-semantic"),
        }

        jobs[job_id].update({
            "status": "done", "progress": 100, "current_step": "Done",
            "matched": matched, "unmatched": unmatched, "skipped": skipped,
            "mappings": _rows_to_flat({"matched": matched, "unmatched": unmatched, "skipped": skipped}),
            "stats": stats,
        })
        _save_job(job_id)   # persist to disk — survives server restarts

    except Exception as exc:
        import traceback
        jobs[job_id].update({
            "status": "error", "progress": 0,
            "current_step": f"Error: {exc}", "error": str(exc),
            "traceback": traceback.format_exc(),
        })


# ── SSE progress ───────────────────────────────────────────────────────────────
import asyncio

@app.get("/api/status/{job_id}")
async def status_stream(job_id: str):
    if job_id not in jobs:
        # Try reloading from disk cache (server may have restarted mid-session)
        _load_caches()
    if job_id not in jobs:
        raise HTTPException(404, "Job not found")

    async def _gen():
        while True:
            job = jobs.get(job_id, {})
            yield f"data: {json.dumps({'progress': job.get('progress', 0), 'step': job.get('current_step', ''), 'status': job.get('status', 'processing'), 'error': job.get('error')})}\n\n"
            if job.get("status") in ("done", "error"):
                break
            await asyncio.sleep(0.3)

    return StreamingResponse(_gen(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ── Download input file ────────────────────────────────────────────────────────
@app.get("/api/download/input/{job_id}")
async def download_input(job_id: str):
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    fp = job.get("file_path", "")
    if not fp or not Path(fp).exists():
        raise HTTPException(404, "Input file not found on server")
    return FileResponse(fp, filename=Path(fp).name)


# ── Get mappings ───────────────────────────────────────────────────────────────
@app.get("/api/mappings/{job_id}")
async def get_mappings(job_id: str):
    job = jobs.get(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    # Always sync from disk if disk cache has more rows (handles stale in-memory state)
    cache_file = CACHE_DIR / f"{job_id}.json"
    if cache_file.exists():
        try:
            disk = json.loads(cache_file.read_text(encoding="utf-8"))
            disk_mappings = disk.get("mappings", [])
            if len(disk_mappings) > len(job.get("mappings", [])):
                job["mappings"] = disk_mappings
                job["stats"]    = disk.get("stats", job.get("stats", {}))
        except Exception:
            pass
    return {
        "mappings": job.get("mappings", []),
        "stats":    job.get("stats", {}),
        "filename": job.get("filename", ""),
    }


# ── Chat ───────────────────────────────────────────────────────────────────────
class ChatBody(BaseModel):
    job_id: str
    message: str
    mappings: list[dict] | None = None   # frontend sends current S.mappings
    reset_history: bool = False          # clear session history before this turn


@app.post("/api/chat")
async def chat(body: ChatBody):
    # Accept chat even if job is no longer in memory (server restart)
    # — the frontend sends its own mappings so the snapshot is always fresh

    if body.reset_history:
        sessions[body.job_id] = []
    history = sessions.setdefault(body.job_id, [])
    history.append({"role": "user", "content": body.message})

    try:
        system_prompt = _build_system(body.job_id, body.mappings)

        resp = claude.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=8192,
            system=system_prompt,
            messages=history,
        )

        raw    = resp.content[0].text
        action = _extract_action(raw)

        # If model promised a change but didn't emit an action block, force a retry
        _UPDATE_INTENT = re.compile(
            r"i'?ll (go ahead|update|change|correct|fix|set|map)|"
            r"(will|going to) (update|change|correct|fix|set|map)|"
            r"update this mapping|let me (update|change|fix)|"
            r"i can (update|change|fix|make)",
            re.IGNORECASE,
        )
        _REFUSAL = re.compile(
            r"don'?t have (visibility|access|information|data)|"
            r"not (listed|included|available|visible)|"
            r"cannot (update|change|access|see|find)|"
            r"not in (the|my) (snapshot|data|list|mapping)|"
            r"only (unmatched|the listed)|"
            r"would need (the|more|additional)|"
            r"unable to (update|change|find|access)",
            re.IGNORECASE,
        )
        user_wants_change = re.search(r"\b(update|change|set|correct|fix|map)\b", body.message, re.IGNORECASE)

        _AWAITING_APPROVAL = re.compile(
            r"shall i apply|should i apply|want me to apply|go ahead and apply|"
            r"apply (all|these|the) \d+ correction|apply (these|the) changes|"
            r"confirm (and i|to apply)|ready to apply|proceed with",
            re.IGNORECASE,
        )
        waiting_for_approval = bool(_AWAITING_APPROVAL.search(raw))

        needs_retry = (not action and not waiting_for_approval
                       and (_UPDATE_INTENT.search(raw) or (user_wants_change and _REFUSAL.search(raw))))
        if needs_retry:
            retry_history = history + [
                {"role": "assistant", "content": raw},
                {"role": "user", "content":
                    "The mapping snapshot I sent contains ALL rows — matched and unmatched. "
                    "You are allowed to update ANY row regardless of its current status. "
                    "Please emit the action block NOW to apply the requested change. "
                    "Use the exact raw_name from the mapping data. Do not explain further."},
            ]
            resp2  = claude.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=512,
                system=system_prompt,
                messages=retry_history,
            )
            raw2   = resp2.content[0].text
            action = _extract_action(raw2)

        display = _strip_action(raw)
        history.append({"role": "assistant", "content": raw})
        _save_chat(body.job_id)
        return {"message": display, "action": action}

    except Exception as exc:
        if history and history[-1]["role"] == "user":
            history.pop()
        import traceback
        traceback.print_exc()
        raise HTTPException(500, detail=f"Chat error: {exc}")


# ── Clear chat history ────────────────────────────────────────────────────────
@app.post("/api/clear-chat/{job_id}")
async def clear_chat(job_id: str):
    sessions[job_id] = []
    (CACHE_DIR / f"{job_id}_chat.json").unlink(missing_ok=True)
    return {"ok": True}


# ── Apply action ───────────────────────────────────────────────────────────────
class ApplyBody(BaseModel):
    job_id: str
    action: dict


@app.post("/api/apply-action")
async def apply_action(body: ApplyBody):
    if body.job_id not in jobs:
        raise HTTPException(404, "Job not found")

    mappings  = jobs[body.job_id]["mappings"]
    id_map    = {r["id"]: i for i, r in enumerate(mappings)}
    name_map  = {r["raw_name"].lower().strip(): i for i, r in enumerate(mappings)}
    affected: list[str] = []

    for change in body.action.get("changes", []):
        field = change.get("field")
        value = change.get("value")
        # Fix model encoding dept change as catalogue_name "Pathology X" / "Radiology X"
        if field == "catalogue_name" and isinstance(value, str):
            if re.match(r'^pathology\s+', value, re.IGNORECASE):
                field, value = "lab_type", "Pathology"
            elif re.match(r'^radiology\s+', value, re.IGNORECASE):
                field, value = "lab_type", "Radiology"
        # Match by raw_name (primary) or row_id (fallback)
        raw_name = (change.get("raw_name") or "").lower().strip()
        rid      = change.get("row_id")
        if raw_name and raw_name in name_map:
            i = name_map[raw_name]
        elif rid and rid in id_map:
            i = id_map[rid]
        else:
            continue
        if field == "catalogue_name":
            mappings[i]["standard_name"]  = value
            if mappings[i].get("status") != "skipped":
                mappings[i]["status"] = "matched" if value else "unmatched"
        elif field == "status":
            mappings[i]["status"] = value
        elif field == "price":
            mappings[i]["price"] = str(value).strip()
        elif field in ("lab_type", "department", "dept"):
            v = str(value).strip()
            if v.lower() in ("radiology", "rad"):
                v = "Radiology"
            elif v.lower() in ("pathology", "path"):
                v = "Pathology"
            mappings[i]["lab_type"] = v
            mappings[i]["_dept"] = v.lower()
        mappings[i]["highlight"] = "ai"
        affected.append(mappings[i]["id"])

    # Sync back to matched/unmatched lists
    jobs[body.job_id]["mappings"] = mappings
    return {"affected": affected, "mappings": mappings}


# ── Update single mapping ──────────────────────────────────────────────────────
class UpdateBody(BaseModel):
    job_id: str
    row_id: str
    standard_name: str | None = None
    price: str | None = None
    lab_type: str | None = None
    keep_status: bool = False


@app.patch("/api/mapping")
async def update_mapping(body: UpdateBody):
    if body.job_id not in jobs:
        raise HTTPException(404, "Job not found")
    for row in jobs[body.job_id]["mappings"]:
        if row["id"] == body.row_id:
            if body.standard_name is not None:
                row["standard_name"] = body.standard_name
                if not body.keep_status and row.get("status") != "skipped":
                    row["status"] = "matched" if body.standard_name.strip() else "unmatched"
                row["highlight"]     = None
            if body.price is not None:
                row["price"] = body.price
            if body.lab_type is not None:
                row["lab_type"] = body.lab_type
            _save_job(body.job_id)
            # Persist correction so future jobs auto-apply it
            if body.standard_name is not None and body.standard_name.strip():
                try:
                    apply_learnings([{
                        "type":               "edited",
                        "provider_name":      row["raw_name"],
                        "old_catalogue_name": row.get("standard_name", ""),
                        "new_catalogue_name": body.standard_name,
                        "lab_type":           row.get("lab_type", ""),
                        "price":              row.get("price", ""),
                    }], PROJECT_ROOT)
                except Exception:
                    pass
            return row
    raise HTTPException(404, "Row not found")


# ── Update row status ──────────────────────────────────────────────────────────
class StatusBody(BaseModel):
    job_id: str
    row_id: str
    status: str   # "matched" | "unmatched" | "skipped"


@app.patch("/api/mapping/status")
async def update_mapping_status(body: StatusBody):
    if body.job_id not in jobs:
        raise HTTPException(404, "Job not found")
    if body.status not in ("matched", "unmatched", "skipped", "confirmed", "rejected"):
        raise HTTPException(400, "Invalid status")
    for row in jobs[body.job_id]["mappings"]:
        if row["id"] == body.row_id:
            row["status"] = body.status
            row["highlight"] = None
            _save_job(body.job_id)
            return row
    raise HTTPException(404, "Row not found")


# ── Master detail lookup ───────────────────────────────────────────────────────
@app.get("/api/master-detail")
async def master_detail(catalogue_name: str):
    """Return Output Format fields from master_details for a matched catalogue name."""
    db_path = _master_db_path()
    if not db_path.exists():
        return {"row": {}}
    try:
        conn = sqlite3.connect(str(db_path))
        # Check if master_details table exists
        tbl = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name='master_details'"
        ).fetchone()
        if not tbl:
            conn.close()
            return {"row": {}}
        cols_info = conn.execute("PRAGMA table_info(master_details)").fetchall()
        col_names = [r[1] for r in cols_info]
        row = conn.execute(
            f"SELECT {', '.join(col_names)} FROM master_details WHERE catalogue_name = ? LIMIT 1",
            (catalogue_name,)
        ).fetchone()
        conn.close()
        if not row:
            return {"row": {}}
        return {"row": dict(zip(col_names, row))}
    except Exception as exc:
        return {"row": {}, "error": str(exc)}


class BulkDetailRequest(BaseModel):
    catalogue_names: list[str]

@app.post("/api/master-details-bulk")
async def master_details_bulk(req: BulkDetailRequest):
    """Return Output Format fields for multiple catalogue names in one call."""
    db_path = _master_db_path()
    if not db_path.exists() or not req.catalogue_names:
        return {"rows": {}}
    try:
        conn = sqlite3.connect(str(db_path))
        tbl = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name='master_details'"
        ).fetchone()
        if not tbl:
            conn.close()
            return {"rows": {}}
        col_names = [r[1] for r in conn.execute("PRAGMA table_info(master_details)").fetchall()]
        placeholders = ",".join("?" * len(req.catalogue_names))
        db_rows = conn.execute(
            f"SELECT {', '.join(col_names)} FROM master_details WHERE catalogue_name IN ({placeholders})",
            req.catalogue_names,
        ).fetchall()
        conn.close()
        result = {row[0]: dict(zip(col_names, row)) for row in db_rows}
        return {"rows": result, "columns": col_names}
    except Exception as exc:
        return {"rows": {}, "error": str(exc)}


# ── Export ─────────────────────────────────────────────────────────────────────
class ExportBody(BaseModel):
    job_id: str
    mappings: list[dict] | None = None   # frontend sends current S.mappings with _masterOverrides


_match_py_module = None  # cached import of match.py


def _load_match_py():
    """Lazily import match.py to reuse its full matching pipeline."""
    global _match_py_module
    if _match_py_module is not None:
        return _match_py_module
    import importlib.util
    candidates = [
        PROJECT_ROOT / ".claude" / "skills" / "process-catalogue" / "scripts" / "match.py",
        PROJECT_ROOT / "console" / ".claude" / "skills" / "process-catalogue" / "scripts" / "match.py",
    ]
    for p in candidates:
        if p.exists():
            spec = importlib.util.spec_from_file_location("_match_py", str(p))
            mod  = importlib.util.module_from_spec(spec)
            try:
                spec.loader.exec_module(mod)
                _match_py_module = mod
                return mod
            except Exception as exc:
                print(f"[WARN] Could not load match.py from {p}: {exc}", file=sys.stderr)
    return None


def _fetch_master_details_map(catalogue_names: list[str]) -> dict[str, dict]:
    """Return dict of catalogue_name → master_details row from SQLite.

    For entries missing important fields, runs the same full matching pipeline
    as match.py (_catalogue_token_match: normalize, typo correction, abbreviation
    expansion, token_sort_ratio A/B/C/1/2 strategies, modality coherence gate)
    to find a better-matching catalogue row and fill the gaps.
    """
    db_path = _master_db_path()
    if not db_path.exists() or not catalogue_names:
        return {}
    try:
        conn = sqlite3.connect(str(db_path))
        tbl = conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name='master_details'"
        ).fetchone()
        if not tbl:
            conn.close()
            return {}
        col_names = [r[1] for r in conn.execute("PRAGMA table_info(master_details)").fetchall()]

        # Exact match first
        placeholders = ",".join("?" * len(catalogue_names))
        rows = conn.execute(
            f"SELECT {', '.join(col_names)} FROM master_details WHERE catalogue_name IN ({placeholders})",
            catalogue_names,
        ).fetchall()
        result = {row[0]: dict(zip(col_names, row)) for row in rows}

        # Fields that should be filled via fuzzy fallback when empty
        FILL_FIELDS = [
            "report_tat", "fasting_hours", "fasting_required", "collection_type",
            "home_collection", "entity_type", "lab_requirement", "age_range",
            "gender", "minimum_patient", "precautions", "description",
            "alias", "tags", "prescription_required", "mrp", "loinc_id",
        ]

        # Rows that still have at least one empty important field
        needs_fill = [
            cname for cname, row in result.items()
            if any(not str(row.get(f, "") or "").strip() for f in FILL_FIELDS)
        ]

        if needs_fill:
            # Load all master_details rows once
            all_rows    = conn.execute(
                f"SELECT {', '.join(col_names)} FROM master_details"
            ).fetchall()
            all_details = [dict(zip(col_names, r)) for r in all_rows]
            all_cnames  = [d["catalogue_name"] for d in all_details]
            name_to_row = {d["catalogue_name"]: d for d in all_details}

            # Try to use match.py's full _catalogue_token_match pipeline
            match_mod = _load_match_py()
            if match_mod is not None:
                # Build norm_cats once (same format _catalogue_token_match expects)
                norm_cats: dict[str, str] = {}
                for c in all_cnames:
                    nc = match_mod.normalize_catalogue(c)
                    if c and nc not in norm_cats:
                        norm_cats[nc] = c

                for cname in needs_fill:
                    primary = result[cname]
                    # Pre-process with same pipeline as match.py
                    corrected  = match_mod.fix_medical_typos(cname)
                    norm_input = match_mod.normalize_catalogue(corrected)

                    # KNOWN_ABBREVIATIONS lookup
                    abbrev_cat = match_mod.KNOWN_ABBREVIATIONS.get(norm_input)
                    if abbrev_cat and abbrev_cat in name_to_row and abbrev_cat != cname:
                        fuzzy_row = name_to_row[abbrev_cat]
                        for field in FILL_FIELDS:
                            if not str(primary.get(field, "") or "").strip():
                                val = fuzzy_row.get(field, "")
                                if str(val or "").strip():
                                    primary[field] = val
                        continue

                    # Full _catalogue_token_match (strategies A/B/C/1/2)
                    match_result = match_mod._catalogue_token_match(
                        norm_input, corrected,
                        None,                # master_df not needed when norm_cats provided
                        all_cat_names=all_cnames,
                        norm_cats=norm_cats,
                    )
                    if match_result:
                        match_name, score = match_result
                        if match_name != cname and score >= 0.65:
                            fuzzy_row = name_to_row.get(match_name, {})
                            for field in FILL_FIELDS:
                                if not str(primary.get(field, "") or "").strip():
                                    val = fuzzy_row.get(field, "")
                                    if str(val or "").strip():
                                        primary[field] = val

            else:
                # Fallback: simple token_sort_ratio ≥ 75 if match.py unavailable
                from rapidfuzz import process as _rp, fuzz as _rf
                for cname in needs_fill:
                    primary = result[cname]
                    matches = _rp.extract(
                        cname, all_cnames, scorer=_rf.token_sort_ratio, limit=5
                    )
                    for match_name, score, _ in matches:
                        if match_name == cname or score < 75:
                            continue
                        fuzzy_row = name_to_row.get(match_name, {})
                        for field in FILL_FIELDS:
                            if not str(primary.get(field, "") or "").strip():
                                val = fuzzy_row.get(field, "")
                                if str(val or "").strip():
                                    primary[field] = val
                        break

        conn.close()
        return result
    except Exception:
        import traceback; traceback.print_exc()
        return {}


class SaveOverridesBody(BaseModel):
    job_id: str
    mappings: list[dict]

@app.post("/api/save-overrides")
async def save_overrides(body: SaveOverridesBody):
    """Persist confirmed _masterOverrides back into the server job cache."""
    job = jobs.get(body.job_id)
    if job:
        # Merge confirmed overrides into server-side mappings
        server_map = {r["id"]: r for r in job.get("mappings", []) if r.get("id")}
        for m in body.mappings:
            mid = m.get("id")
            if mid and mid in server_map and m.get("_masterOverrides"):
                server_map[mid]["_masterOverrides"] = m["_masterOverrides"]
        job["mappings"] = list(server_map.values())
        _save_job(body.job_id)
    return {"ok": True}


@app.post("/api/export")
async def export(body: ExportBody):
    job = jobs.get(body.job_id)
    if not job:
        raise HTTPException(404, "Job not found")

    # Prefer frontend mappings (include _masterOverrides); fall back to server copy
    mappings = body.mappings or job["mappings"]

    # Fetch master details for all matched catalogue names
    catalogue_names = list({r.get("standard_name","") for r in mappings if r.get("standard_name")})
    master_detail_map = _fetch_master_details_map(catalogue_names)

    def _build_row(r: dict, status: str) -> dict:
        cat_name = r.get("standard_name", "") or ""
        md = master_detail_map.get(cat_name, {})
        overrides = r.get("_masterOverrides") or {}
        # Merge: overrides win over master_details
        merged = {**md, **overrides}
        dept = r.get("lab_type", "")
        if dept and dept.lower() not in ("pathology", "radiology"):
            dept = ""
        return {
            "provider_name":   r.get("raw_name", ""),
            "catalogue_name":  cat_name,
            "match_type":      r.get("match_type", status),
            "confidence":      r.get("confidence", 0),
            "price":           r.get("price", ""),
            "department":      dept,
            # Master detail fields
            "loinc_id":              merged.get("loinc_id", ""),
            "lab_requirement":       merged.get("lab_requirement", ""),
            "entity_type":           merged.get("entity_type", ""),
            "mrp":                   merged.get("mrp", "") or r.get("price", ""),
            "discounted_price":      merged.get("discounted_price", ""),
            "display_mrp":           merged.get("display_mrp", "") or r.get("price", ""),
            "collection_type":       merged.get("collection_type", ""),
            "home_collection":       merged.get("home_collection", ""),
            "fasting_required":      merged.get("fasting_required", ""),
            "fasting_hours":         merged.get("fasting_hours", ""),
            "age_range":             merged.get("age_range", "").replace("$", " to "),
            "gender":                merged.get("gender", ""),
            "minimum_patient":       merged.get("minimum_patient", ""),
            "report_tat":            merged.get("report_tat", ""),
            "alias":                 merged.get("alias", ""),
            "tags":                  merged.get("tags", ""),
            "prescription_required": merged.get("prescription_required", ""),
            "precautions":           merged.get("precautions", ""),
            "description":           merged.get("description", ""),
            "provider_item_name":    re.sub(r'[^a-z0-9]+', '-', cat_name.strip().lower()).strip('-') if cat_name.strip() else '',
        }

    matched_rows   = [_build_row(r, "MATCHED")
                      for r in mappings if r.get("status") in ("matched", "confirmed") and r.get("standard_name")]
    unmatched_rows = [_build_row(r, "UNMATCHED")
                      for r in mappings if r.get("status") == "unmatched"]

    corrections = [{"type": "edited", "provider_name": r["raw_name"],
                    "old_catalogue_name": "", "new_catalogue_name": r.get("standard_name","")}
                   for r in mappings if r.get("status") == "matched" and r.get("highlight") == "ai"]

    out_path = generate_output_excel(body.job_id, matched_rows, unmatched_rows, job["filename"], PROJECT_ROOT)

    if corrections:
        apply_learnings(corrections, PROJECT_ROOT)

    return FileResponse(out_path, filename=Path(out_path).name,
                        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ── Static (must be last) ──────────────────────────────────────────────────────
app.mount("/", StaticFiles(directory=str(FRONTEND_DIR), html=True), name="static")
