"""
Provider Catalogue Reconciliation — Flask server.

All processing logic is delegated to the Claude Code CLI (process-catalogue skill).
This file handles: file upload, session management, skill invocation,
output parsing, user corrections, learning, and Excel download.

Run:
    python3 server.py
"""

from __future__ import annotations

import csv
import json
import os
import re
import subprocess
import sys
import tempfile
import threading
import uuid
import webbrowser
from datetime import datetime
from pathlib import Path

# ── Bootstrap Flask ───────────────────────────────────────────────────────────
try:
    from flask import Flask, jsonify, request, send_file, send_from_directory
except ImportError:
    subprocess.run([sys.executable, "-m", "pip", "install", "-q", "flask"], check=True)
    from flask import Flask, jsonify, request, send_file, send_from_directory

# ── Bootstrap minimal deps (Excel output generation only) ────────────────────
def _ensure(*pkgs: str) -> None:
    missing = [p for p in pkgs if not _importable(p)]
    if missing:
        subprocess.run([sys.executable, "-m", "pip", "install", "-q"] + missing, check=False)

def _importable(pkg: str) -> bool:
    try:
        __import__(pkg.split("[")[0].replace("-", "_"))
        return True
    except ImportError:
        return False

_ensure("pandas", "openpyxl")
_ensure("pdfplumber")
_ensure("python-docx")
_ensure("docx2txt")
_ensure("Pillow", "pytesseract")
_ensure("rapidfuzz")

import pandas as pd

# ── Paths ─────────────────────────────────────────────────────────────────────
HERE         = Path(__file__).parent
PROJECT_ROOT = HERE.parent
OUTPUT_DIR   = PROJECT_ROOT / "output"
INPUT_DIR    = PROJECT_ROOT / "input"
LEARNING_DIR = PROJECT_ROOT / "learning"
MATCH_PY     = PROJECT_ROOT / ".claude" / "skills" / "process-catalogue" / "scripts" / "match.py"

OUTPUT_DIR.mkdir(exist_ok=True)
INPUT_DIR.mkdir(exist_ok=True)
LEARNING_DIR.mkdir(exist_ok=True)

# ── Flask app ─────────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder=str(HERE), static_url_path="")

@app.route("/")
def index():
    return send_from_directory(str(HERE), "index.html")


# ═══════════════════════════════════════════════════════════════════
# Session store  (in-memory, keyed by uuid)
# ═══════════════════════════════════════════════════════════════════
_SESSIONS: dict[str, dict] = {}


# ═══════════════════════════════════════════════════════════════════
# /api/process  —  save file, start Claude skill in background
# ═══════════════════════════════════════════════════════════════════
@app.route("/api/process", methods=["POST"])
def process():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "No file provided"}), 400

    filename  = file.filename or "upload"
    safe_name = re.sub(r"[^\w.\-]", "_", filename)
    input_path = INPUT_DIR / safe_name
    file.save(str(input_path))

    sid = str(uuid.uuid4())
    _SESSIONS[sid] = {
        "status":   "queued",
        "filename": filename,
        "step":     0,
        "error":    None,
        "result":   None,
    }

    threading.Thread(
        target=_run_skill,
        args=(sid, input_path, filename),
        daemon=True,
    ).start()

    return jsonify({"session_id": sid, "filename": filename})


# ═══════════════════════════════════════════════════════════════════
# /api/status/<session_id>  —  poll for skill progress / result
# ═══════════════════════════════════════════════════════════════════
@app.route("/api/status/<session_id>")
def status(session_id: str):
    s = _SESSIONS.get(session_id)
    if not s:
        return jsonify({"error": "Session not found"}), 404
    return jsonify({
        "status":   s["status"],
        "step":     s["step"],
        "filename": s["filename"],
        "error":    s["error"],
        "result":   s["result"] if s["status"] == "done" else None,
    })


# ═══════════════════════════════════════════════════════════════════
# Skill runner  (background thread)
# ═══════════════════════════════════════════════════════════════════
# ═══════════════════════════════════════════════════════════════════
# Pipeline runner  (background thread — no network / API required)
# Follows logic defined in:
#   .claude/skills/process-catalogue/references/parsing-guide.md
#   .claude/skills/process-catalogue/references/accuracy-loop-guide.md
#   .claude/skills/process-catalogue/scripts/match.py
# ═══════════════════════════════════════════════════════════════════
def _run_skill(sid: str, input_path: Path, filename: str) -> None:
    s = _SESSIONS[sid]
    s["status"] = "running"
    try:
        # ── Step 1: Parse file ────────────────────────────────────────────────
        s["step"] = 1
        names = _parse_file(str(input_path))
        if not names:
            s["status"] = "error"
            s["error"]  = "No test names found in the file"
            return

        # ── Step 2: Deterministic matching via match.py ───────────────────────
        s["step"] = 2
        tmp_dir = Path(tempfile.mkdtemp(prefix="catalogue_"))
        results = _run_match(names, tmp_dir)

        matched:   list[dict] = []
        unmatched: list[dict] = []
        skipped:   list[dict] = []

        for row in results:
            mt   = str(row.get("Match Type", "UNMATCHED"))
            item = {
                "id":            str(uuid.uuid4()),
                "provider_name": str(row.get("Provider Test Name", "")),
                "catalogue_name":str(row.get("Catalogue Test Name") or ""),
                "match_type":    mt,
                "confidence":    float(row.get("Confidence Score") or 0),
                "user_edited":   False,
            }
            if mt == "SKIPPED":
                skipped.append(item)
            elif mt == "UNMATCHED":
                item["suggested_name"] = ""
                unmatched.append(item)
            else:
                matched.append(item)

        # ── Step 3: Semantic / accuracy-loop pass ─────────────────────────────
        s["step"] = 3
        semantic_count  = 0
        semantic_status = "skipped — no unmatched items after fuzzy match"
        if unmatched:
            sem_results, semantic_status = _semantic_pass(unmatched)
            still_unmatched: list[dict] = []
            for item in unmatched:
                key = item["provider_name"].lower().strip()
                sr  = sem_results.get(key)
                if sr and sr.get("catalogue_name") and sr["match_type"] not in ("UNMATCHED", "SKIPPED"):
                    item.update(catalogue_name=sr["catalogue_name"],
                                match_type=sr["match_type"],
                                confidence=sr["confidence"])
                    matched.append(item)
                    semantic_count += 1
                elif sr and sr.get("match_type") == "SKIPPED":
                    item["match_type"] = "SKIPPED"
                    skipped.append(item)
                else:
                    still_unmatched.append(item)
            unmatched = still_unmatched

        # ── Step 4: Done ──────────────────────────────────────────────────────
        s["step"] = 4
        stats = {
            "total":           len(names),
            "matched":         len(matched),
            "unmatched":       len(unmatched),
            "skipped":         len(skipped),
            "exact":           sum(1 for r in results if r.get("Match Type") == "exact"),
            "fuzzy":           sum(1 for r in results if r.get("Match Type") in ("fuzzy", "fuzzy-catalogue")),
            "fuzzy_catalogue": sum(1 for r in results if r.get("Match Type") == "fuzzy-catalogue"),
            "semantic":        semantic_count,
        }
        # Add suggested_name to every skipped row so the frontend can edit it
        for item in skipped:
            item.setdefault("suggested_name", "")

        s["status"] = "done"
        s["result"] = {
            "filename":        filename,
            "output_file":     "",
            "matched":         matched,
            "unmatched":       unmatched,
            "skipped":         skipped,
            "stats":           stats,
            "semantic_status": semantic_status,
        }
    except Exception as exc:
        import traceback
        s["status"] = "error"
        s["error"]  = f"{type(exc).__name__}: {exc}\n{traceback.format_exc()[-2000:]}"


# ═══════════════════════════════════════════════════════════════════
# File parsing  (parsing-guide.md)
# ═══════════════════════════════════════════════════════════════════
_NAME_KWS = {"test","name","item","procedure","service","description",
             "particular","investigation","examination","profile"}
_SKIP_RE  = re.compile(
    r"^(total|subtotal|page\s*\d+|s\.?n\.?o\.?|sr\.?\s*no|sl\.?\s*no|"
    r"provider\s+name|powered\s+by|price|rate|amount|code|"
    r"category|department|section|group)$", re.IGNORECASE)

def _is_valid(v: str) -> bool:
    v = v.strip()
    return bool(v) and 2 <= len(v) <= 200 and not v.replace(" ","").isdigit() and not _SKIP_RE.match(v)

def _dedup(names: list[str]) -> list[str]:
    seen: set[str] = set()
    out:  list[str] = []
    for n in names:
        k = n.lower().strip()
        if k not in seen:
            seen.add(k); out.append(n)
    return out

def _find_col(df) -> int:
    for ri in range(min(10, len(df))):
        for ci, val in enumerate(df.iloc[ri]):
            if isinstance(val, str) and any(k in val.lower() for k in _NAME_KWS):
                return ci
    return 0

def _parse_file(path: str) -> list[str]:
    ext = Path(path).suffix.lower()
    if ext in (".xlsx", ".xls"):   return _from_excel(path)
    if ext == ".pdf":              return _from_pdf(path)
    if ext == ".docx":             return _from_docx(path)
    if ext == ".doc":              return _from_doc(path)
    if ext == ".csv":              return _from_csv(path)
    if ext in (".png",".jpg",".jpeg",".tiff",".bmp",".webp"): return _from_image(path)
    raise ValueError(f"Unsupported file type: {ext!r}")

def _from_excel(path: str) -> list[str]:
    xl   = pd.ExcelFile(path)
    best: list[str] = []
    for sheet in xl.sheet_names:
        raw = pd.read_excel(path, sheet_name=sheet, header=None, dtype=str)
        if raw.empty: continue
        ci = _find_col(raw)
        hr = 0
        for r in range(min(10, len(raw))):
            v = str(raw.iloc[r, ci] or "").strip()
            if v and any(k in v.lower() for k in _NAME_KWS):
                hr = r + 1; break
        names = [n.strip() for n in raw.iloc[hr:, ci].dropna().astype(str) if _is_valid(n.strip())]
        if len(names) > len(best): best = names
    if not best:
        df = pd.read_excel(path, dtype=str)
        best = [n.strip() for n in df.iloc[:,0].dropna().astype(str) if _is_valid(n.strip())]
    return _dedup(best)

def _from_pdf(path: str) -> list[str]:
    try:
        import pdfplumber
        names: list[str] = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_names: list[str] = []
                for tbl in (page.extract_tables() or []):
                    for row in tbl:
                        for cell in (row or []):
                            v = str(cell or "").strip()
                            if _is_valid(v): page_names.append(v)
                if not page_names:
                    for line in (page.extract_text() or "").split("\n"):
                        if _is_valid(line.strip()): page_names.append(line.strip())
                names.extend(page_names)
        return _dedup(names)
    except ImportError:
        pass
    try:
        import fitz
        names = []
        for page in fitz.open(path):
            for line in (page.get_text() or "").split("\n"):
                if _is_valid(line.strip()): names.append(line.strip())
        return _dedup(names)
    except ImportError:
        raise RuntimeError("pdfplumber not installed — run: pip install pdfplumber")

def _from_docx(path: str) -> list[str]:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("python-docx not installed — run: pip install python-docx")
    doc   = Document(path)
    names: list[str] = []
    for tbl in doc.tables:
        tc = None
        if tbl.rows:
            for i, h in enumerate(c.text.strip() for c in tbl.rows[0].cells):
                if any(k in h.lower() for k in _NAME_KWS): tc = i; break
            if tc is None: tc = 0
        for ri, row in enumerate(tbl.rows):
            if ri == 0: continue
            if tc is not None and tc < len(row.cells):
                t = row.cells[tc].text.strip()
                if _is_valid(t): names.append(t)
    if not names:
        for p in doc.paragraphs:
            if _is_valid(p.text.strip()): names.append(p.text.strip())
    return _dedup(names)

def _from_doc(path: str) -> list[str]:
    try:
        import docx2txt
        text  = docx2txt.process(path)
        names = [ln.strip() for ln in (text or "").split("\n") if _is_valid(ln.strip())]
        if names: return _dedup(names)
    except (ImportError, Exception):
        pass
    try:
        return _from_docx(path)
    except Exception:
        pass
    raise RuntimeError("Cannot read .doc file. Save as .docx and retry, or: pip install docx2txt")

def _from_image(path: str) -> list[str]:
    try:
        from PIL import Image, ImageFilter, ImageEnhance
    except ImportError:
        raise RuntimeError("Pillow not installed — run: pip install Pillow")
    try:
        import pytesseract
    except ImportError:
        raise RuntimeError("pytesseract not installed — run: pip install pytesseract")
    img = Image.open(path).convert("L")
    img = ImageEnhance.Contrast(img).enhance(2.0)
    img = img.filter(ImageFilter.SHARPEN)
    try:
        raw = pytesseract.image_to_string(img, config="--psm 6 --oem 3")
    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "Tesseract OCR not found.\n"
            "  Mac:   brew install tesseract\n"
            "  Linux: sudo apt-get install tesseract-ocr"
        )
    names: list[str] = []
    for line in raw.split("\n"):
        line = re.sub(r"[|\\]{2,}", " ", line)
        line = re.sub(r"\s+", " ", line).strip()
        if _is_valid(line): names.append(line)
    if not names:
        raise RuntimeError("OCR ran but found no valid test names. Try a higher-resolution scan.")
    return _dedup(names)

def _from_csv(path: str) -> list[str]:
    df = pd.read_csv(path, dtype=str)
    for col in df.columns:
        if any(k in col.lower() for k in _NAME_KWS):
            return _dedup([n.strip() for n in df[col].dropna().astype(str) if _is_valid(n.strip())])
    return _dedup([n.strip() for n in df.iloc[:,0].dropna().astype(str) if _is_valid(n.strip())])


# ═══════════════════════════════════════════════════════════════════
# match.py runner  (matching-guide.md)
# ═══════════════════════════════════════════════════════════════════
MASTER_FILE = PROJECT_ROOT / "refrences" / "master_file.xlsx"

def _run_match(names: list[str], tmp_dir: Path) -> list[dict]:
    input_csv  = tmp_dir / "extracted_names.csv"
    output_csv = tmp_dir / "matched_results.csv"
    with open(input_csv, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["Provider Test Name"])
        for n in names: w.writerow([n])
    result = subprocess.run(
        [sys.executable, str(MATCH_PY),
         "--input",  str(input_csv),
         "--master", str(MASTER_FILE),
         "--output", str(output_csv)],
        capture_output=True, text=True, timeout=180, cwd=str(PROJECT_ROOT),
    )
    if result.returncode != 0:
        raise RuntimeError(f"match.py failed:\n{result.stderr[-2000:]}")
    return pd.read_csv(output_csv).to_dict("records")


# ═══════════════════════════════════════════════════════════════════
# Semantic pass  (accuracy-loop-guide.md — pure Python, no network)
# ═══════════════════════════════════════════════════════════════════
_SEM_ABBREVS: dict[str, str] = {
    "trbc":"Total RBC", "hvc":"HCV", "ckmb":"CK-MB", "c k m b":"CK-MB",
    "ck mb":"CK-MB", "e 2":"Estradiol", "lactac":"Lactate", "zsr":"ESR",
    "lft":"Liver Function Test", "kft":"Kidney Function Test",
    "rft":"Kidney Function Test", "hba1c":"Glycated Haemoglobin",
    "abpa":"Aspergillus Fumigatus Antibodies IgG",
    "uds":"Drugs of Abuse, 7 Drugs Urine Screen", "pdger":"PDGFRA",
    "t zunck":"Tzanck Smear",
}

_SEM_TYPOS: list[tuple[str, str]] = [
    (r"\bpsychological\b","physiological"), (r"\bphysilogical\b","physiological"),
    (r"\bbroanchoscopy\b","bronchoscopy"),  (r"\btynconametr[yi]\b","tympanometry"),
    (r"\bt\s+zunck\b","tzanck smear"),      (r"\bgene\s+xpart\b","genexpert"),
    (r"\burine\s+drug\s+essay\b","urine drug screen"),
    (r"\bpdger\b","pdgfra"),                (r"\bimmunoflorescence\b","immunofluorescence"),
    (r"\bqualitaitve\b","qualitative"),     (r"\badreanal\b","adrenal"),
]

_SEM_FALLBACKS: list[tuple[str, str]] = [
    (r"\b(sugar|glucose)\b.*(csf|cerebrospinal)|(csf|cerebrospinal).*\b(sugar|glucose)\b","Glucose for CSF"),
    (r"\bada\b.*(csf|cerebrospinal)|(csf|cerebrospinal).*\bada\b","Adenosine Deaminase (ADA), CSF"),
    (r"\bpandy\b","Protein CSF"),
    (r"\bbiochem\w*.*(csf|cerebrospinal)|(csf|cerebrospinal).*\bbiochem","Fluid Examination Biochemistry"),
    (r"\bglucose\b.*(synovial|periton|ascit|drain|body\s*fluid)|(synovial|periton|ascit|drain).*\bglucose\b","GLUCOSE, BODY FLUID"),
    (r"\bcalcium\b.*(body\s*fluid|drain|synovial|periton)|(body\s*fluid|drain).*\bcalcium\b","Calcium, Body fluids"),
    (r"\bham\s+test\b|\bauto.*hemolys","HAM Test (Acidified Lysis Test)"),
    (r"\bantenatal.*antibody.*screen","Blood group Unexpected antibody screen, Blood"),
    (r"\bzeta\s+sedimentation","ESR"),
    (r"\bhplc\b.*(hb|haemoglobin|hemoglobin)","Abnormal Haemoglobin Studies(Hb Variant), Blood"),
    (r"\bacute\s+lymphoblastic\s+leuk","Leukemia-Acute Panel By Flowcytometry, Blood"),
    (r"\bpdg[ef][a-z]*\s*.*mutation|\bmutation.*pdg[ef][a-z]*","PDGFRA Mutation Analysis in blocks"),
    (r"\bimmuno\w*flu[ao]r.*malaria","MALARIAL PARASITE (FLUORESCENCE BASED DETECTION)"),
    (r"\bstool.*fat\b|\bfecal.*fat\b|\bfat.*stool\b","Stool For Fat Globules (Sudan IV Stain)"),
    (r"\burine.*fat\b|\bfat.*urine\b","Urine For Fat Globules"),
    (r"\bsterco[a-z]*linog","Urobilinogen Random Urine"),
    (r"\brenal\s+panel\b","Kidney Function Test"),
    (r"\bcholesterol.*hdl.*ratio|hdl\s*chol.*ratio","Total Cholesterol/ HDL Ratio"),
    (r"\btrh\b.*stimulation","TRH (THYROID RELEASING HORMONE Stimulation test for Prolactin)"),
    (r"\bprolonged.*hypoglycae|\bhypoglycaemic.*test","PROLONGED GTT"),
    (r"\bbone\s+specific\s+alkaline\s+phosphatase","ALK. PHOSPHATASEBONE:Immunoassay (IA)"),
    (r"\bblood\s+spot.*amino|\bamino.*blood\s+spot","Amino Acid Quantitative, Plasma"),
    (r"\b(uterus|endometri)\w*.*biopsy|biopsy.*(uterus|endometri)","Endometrial Biopsy"),
    (r"\bbiopsy\b.*\b(tb|pcr)\b|\b(tb|pcr)\b.*\bbiopsy","TB PCR (DNA) MTB, Body Fluid"),
    (r"\bbiopsy\b","SMALL BIOPSY"),
    (r"\btympanometr","AUDIOMETRY"),
    (r"\busg\b.*soft\s*tissue|\bsoft\s*tissue\b.*usg","USG SOFT TISSUE"),
    (r"\bcolou?r\s*doppler\b.*soft|soft.*\bcolou?r\s*doppler\b","Doppler Soft Part NA Study"),
    (r"\babpa\b","Aspergillus Fumigatus Antibodies IgG"),
    (r"\b(food|drug)\b.*allergy|\ballergy\b.*(food|drug)","Allergy Veg & Non-Veg Panel By Elisa Method"),
    (r"\bdrug\s*(screen|scrn)\b|\buds\b","Drugs of Abuse, 7 Drugs Urine Screen"),
    (r"\burine\b.*(potassium|chloride).*(potassium|chloride)","Electrolyte, 24 Hrs Urine"),
    (r"\busg\b","USG OTHER SPECIFIC REGION"),
]

_COMBO_RE = re.compile(r"(?<!\()\s+(?:&|and|with)\s+(?![^(]*\))", re.IGNORECASE)

def _sem_norm(name: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9\s]", " ", name.lower())).strip()

def _sem_strip(name: str) -> str:
    name = re.sub(r"\s*[\(\[](special|fasting|fbs|pp|random|emerg\w*|stat|urgent|repeat)[\)\]]",
                  "", name, flags=re.IGNORECASE)
    return re.sub(r"\s*-\s*(?:f|pp|r|s|fbs)\s*$", "", name, flags=re.IGNORECASE).strip()

def _load_catalogue_names() -> list[str]:
    db = PROJECT_ROOT / "refrences" / "master_file.xlsx.db"
    try:
        import sqlite3
        conn = sqlite3.connect(str(db))
        rows = conn.execute("SELECT DISTINCT catalogue_name FROM master ORDER BY id").fetchall()
        conn.close()
        return [r[0] for r in rows if r[0]]
    except Exception:
        try:
            df = pd.read_excel(MASTER_FILE, dtype=str)
            for col in df.columns:
                if any(k in col.lower() for k in ("package","catalogue","test name","item")):
                    return df[col].dropna().astype(str).str.strip().unique().tolist()
            return df.iloc[:,1].dropna().astype(str).str.strip().unique().tolist()
        except Exception:
            return []

def _semantic_pass(unmatched_items: list[dict]) -> tuple[dict[str, dict], str]:
    if not unmatched_items:
        return {}, "skipped — no unmatched items"
    try:
        from rapidfuzz import fuzz, process as fzp
    except ImportError:
        return {}, "skipped — rapidfuzz not installed"
    catalogue_names = _load_catalogue_names()
    if not catalogue_names:
        return {}, "skipped — catalogue database not available"
    cat_lower = [c.lower() for c in catalogue_names]

    result: dict[str, dict] = {}
    resolved = 0

    for item in unmatched_items:
        orig = item["provider_name"]
        key  = orig.lower().strip()

        # 1. Combination test → SKIPPED
        if _COMBO_RE.search(re.sub(r"\([^)]*\)", "", orig)):
            result[key] = {"catalogue_name":"","match_type":"SKIPPED","confidence":0.0}
            continue

        # 2. Typo correction
        corrected = orig
        for pat, repl in _SEM_TYPOS:
            corrected = re.sub(pat, repl, corrected, flags=re.IGNORECASE)

        # 3. Suffix strip
        corrected = _sem_strip(corrected)

        # 4. Abbreviation expansion
        norm = _sem_norm(corrected)
        if norm in _SEM_ABBREVS:
            corrected = _SEM_ABBREVS[norm]
            norm = _sem_norm(corrected)

        # 5. Fallback patterns
        cat_match: str | None = None
        for pat, cat in _SEM_FALLBACKS:
            if re.search(pat, corrected, re.IGNORECASE):
                cat_match = cat; break

        if cat_match:
            result[key] = {"catalogue_name":cat_match,"match_type":"fuzzy-semantic","confidence":0.90}
            resolved += 1; continue

        # 6. Substring search
        norm_words = [w for w in norm.split() if len(w) > 2]
        if len(norm_words) >= 2:
            for ci, cname in enumerate(cat_lower):
                if all(w in cname for w in norm_words):
                    result[key] = {"catalogue_name":catalogue_names[ci],"match_type":"fuzzy-semantic","confidence":0.80}
                    resolved += 1; cat_match = catalogue_names[ci]; break

        if cat_match: continue

        # 7. Fuzzy token_sort_ratio ≥ 65%
        best = fzp.extractOne(norm, cat_lower, scorer=fuzz.token_sort_ratio, score_cutoff=65.0)
        if best:
            result[key] = {"catalogue_name":catalogue_names[best[2]],"match_type":"fuzzy-semantic","confidence":round(best[1]/100,3)}
            resolved += 1; continue

        result[key] = {"catalogue_name":"","match_type":"UNMATCHED","confidence":0.0}

    return result, f"ran — {resolved} of {len(unmatched_items)} resolved"


# ═══════════════════════════════════════════════════════════════════
# /api/sync  —  persist user corrections without re-generating Excel
# ═══════════════════════════════════════════════════════════════════
@app.route("/api/sync", methods=["POST"])
def sync():
    body        = request.get_json(force=True) or {}
    corrections = body.get("corrections", [])
    try:
        summary = _apply_learnings(corrections) if corrections else []
        return jsonify({"synced": len(corrections), "learning_summary": summary})
    except Exception as exc:
        import traceback
        return jsonify({"error": str(exc), "detail": traceback.format_exc()}), 500


# ═══════════════════════════════════════════════════════════════════
# /api/finalize  —  write user-edited rows to Excel + apply learnings
# ═══════════════════════════════════════════════════════════════════
@app.route("/api/finalize", methods=["POST"])
def finalize():
    body        = request.get_json(force=True) or {}
    matched     = body.get("matched",     [])
    unmatched   = body.get("unmatched",   [])
    corrections = body.get("corrections", [])
    filename    = body.get("filename",    "output")
    try:
        out_path = _generate_excel(matched, unmatched, filename)
        summary  = _apply_learnings(corrections) if corrections else []
        return jsonify({
            "output_filename":  os.path.basename(out_path),
            "learning_summary": summary,
        })
    except Exception as exc:
        import traceback
        return jsonify({"error": str(exc), "detail": traceback.format_exc()}), 500


# ═══════════════════════════════════════════════════════════════════
# /api/download/<filename>
# ═══════════════════════════════════════════════════════════════════
@app.route("/api/download/<path:filename>")
def download(filename: str):
    safe = re.sub(r"[^a-zA-Z0-9_\-\. ]", "", filename)
    path = OUTPUT_DIR / safe
    if not path.exists():
        return jsonify({"error": "File not found"}), 404
    return send_file(
        str(path),
        as_attachment=True,
        download_name=safe,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# ═══════════════════════════════════════════════════════════════════
# Excel generation  (for /api/finalize — writes user-edited rows)
# ═══════════════════════════════════════════════════════════════════
def _read_template_cols() -> list[str]:
    try:
        return list(pd.read_excel(PROJECT_ROOT / "refrences" / "Output_format.xlsx", nrows=0).columns)
    except Exception:
        return []


def _generate_excel(matched: list[dict], unmatched: list[dict], filename: str) -> str:
    import openpyxl
    from openpyxl.styles import Alignment, Font, PatternFill

    tpl = _read_template_cols()

    def _pick(candidates: list[str], default: str) -> str:
        for c in tpl:
            if c.lower().strip() in [x.lower() for x in candidates]:
                return c
        return default

    c_cat  = _pick(["test name", "item name", "catalogue test name", "standard name"], "Catalogue Test Name")
    c_prov = _pick(["provider name", "original name", "provider test name"],            "Provider Test Name")
    c_mt   = _pick(["match type", "type"],                                              "Match Type")
    c_conf = _pick(["confidence", "score", "confidence score"],                         "Confidence Score")
    cols   = [c_cat, c_prov, c_mt, c_conf]

    def _rec(row: dict, status: str) -> dict:
        conf = row.get("confidence", 0) or 0
        return {
            c_cat:  row.get("catalogue_name", ""),
            c_prov: row.get("provider_name",  ""),
            c_mt:   row.get("match_type",     status),
            c_conf: f"{conf:.0%}" if conf else "",
        }

    stem = re.sub(r"[^a-zA-Z0-9_\- ]", "_", Path(filename).stem)[:50].strip("_")
    out  = OUTPUT_DIR / f"{stem}_reconciled.xlsx"
    if out.exists():
        out = OUTPUT_DIR / f"{stem}_reconciled_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"

    wb  = openpyxl.Workbook()
    hf  = Font(color="FFFFFF", bold=True)
    ha  = Alignment(horizontal="center")

    def _sheet(ws, rows: list[dict], status: str, fill_hex: str) -> None:
        fill = PatternFill("solid", fgColor=fill_hex)
        for ci, c in enumerate(cols, 1):
            cell = ws.cell(row=1, column=ci, value=c)
            cell.fill, cell.font, cell.alignment = fill, hf, ha
        for ri, row in enumerate(rows, 2):
            rec = _rec(row, status)
            for ci, c in enumerate(cols, 1):
                ws.cell(row=ri, column=ci, value=rec.get(c, ""))
        for col_cells in ws.columns:
            w = max((len(str(c.value or "")) for c in col_cells), default=12)
            ws.column_dimensions[col_cells[0].column_letter].width = min(w + 4, 55)

    ws_m       = wb.active
    ws_m.title = "Matched"
    _sheet(ws_m, matched,   "MATCHED",   "1E3A5F")
    _sheet(wb.create_sheet("UNMATCHED"), unmatched, "UNMATCHED", "7F1D1D")

    wb.save(str(out))
    return str(out)


# ═══════════════════════════════════════════════════════════════════
# Learning system  —  persist corrections, patch match.py
# ═══════════════════════════════════════════════════════════════════
def _corrections_path() -> Path:
    return LEARNING_DIR / "corrections.json"

def _load_corrections() -> list[dict]:
    p = _corrections_path()
    try:
        return json.loads(p.read_text(encoding="utf-8")) if p.exists() else []
    except Exception:
        return []

def _normalize_key(name: str) -> str:
    name = name.lower().strip()
    name = re.sub(r"[^a-z0-9\s]", "", name)
    return re.sub(r"\s+", " ", name).strip()

def _apply_learnings(corrections: list[dict]) -> list[str]:
    existing    = _load_corrections()
    summary:    list[str]       = []
    new_abbrevs: dict[str, str] = {}

    for c in corrections:
        c["timestamp"] = datetime.now().isoformat()
        existing.append(c)
        if c.get("type") == "approved" and c.get("new_catalogue_name"):
            norm = _normalize_key(c["provider_name"])
            cat  = c["new_catalogue_name"]
            if _abbrev_candidate(norm, cat):
                new_abbrevs[norm] = cat

    _corrections_path().write_text(
        json.dumps(existing, indent=2, default=str), encoding="utf-8"
    )
    summary.append(f"Saved {len(corrections)} correction(s) to learning/corrections.json")

    if new_abbrevs:
        added = _patch_match_py(new_abbrevs)
        for k, v in added.items():
            summary.append(f'Learned: "{k}" → "{v}" (added to match.py KNOWN_ABBREVIATIONS)')

    return summary

def _abbrev_candidate(norm: str, catalogue: str) -> bool:
    if not norm or not catalogue or len(norm) > 15 or norm.isdigit():
        return False
    if not MATCH_PY.exists():
        return False
    content = MATCH_PY.read_text(encoding="utf-8")
    return f'"{norm}"' not in content and f"'{norm}'" not in content

def _patch_match_py(new_abbrevs: dict[str, str]) -> dict[str, str]:
    if not MATCH_PY.exists():
        return {}
    lines  = MATCH_PY.read_text(encoding="utf-8").split("\n")
    ins_at = -1
    in_blk = False
    for i, l in enumerate(lines):
        if "KNOWN_ABBREVIATIONS" in l and "dict[str, str]" in l:
            in_blk = True
        elif in_blk and l.strip() == "}":
            ins_at = i
            break
    if ins_at < 0:
        return {}
    added: dict[str, str] = {}
    content = "\n".join(lines)
    for norm, cat in new_abbrevs.items():
        if f'"{norm}"' in content:
            continue
        pad = " " * max(0, 8 - len(norm))
        lines.insert(ins_at, f'    "{norm}":{pad}"{cat}",  # learned from user')
        ins_at += 1
        added[norm] = cat
    if added:
        MATCH_PY.write_text("\n".join(lines), encoding="utf-8")
    return added


# ═══════════════════════════════════════════════════════════════════
# Entry point
# ═══════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    print(f"\n  ✦  Catalogue Reconciliation  →  http://localhost:{port}")
    print("  ✓  Pipeline: match.py + semantic pass (no API key required)\n")
    threading.Timer(1.0, lambda: webbrowser.open(f"http://localhost:{port}")).start()
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
