"""
File parsing, match.py orchestration, and Excel output generation.
"""

from __future__ import annotations

import csv
import os
import re
import subprocess
import sys
import tempfile
import uuid
from pathlib import Path
from typing import Optional

import pandas as pd


# ── progress helpers ──────────────────────────────────────────────────────────

def _progress(jobs: dict, job_id: str, pct: int, step: str) -> None:
    if job_id in jobs:
        jobs[job_id]["progress"] = pct
        jobs[job_id]["current_step"] = step


# ── dependency bootstrap ──────────────────────────────────────────────────────

def _ensure_deps() -> None:
    missing = []
    for pkg in ("rapidfuzz", "openpyxl", "pdfplumber", "docx"):
        try:
            __import__(pkg)
        except ImportError:
            missing.append(pkg)
    if missing:
        subprocess.run(
            [sys.executable, "-m", "pip", "install", "-q",
             "rapidfuzz", "pandas", "openpyxl", "numpy", "pdfplumber", "python-docx"],
            check=False,
        )


# ── file parsing ──────────────────────────────────────────────────────────────

_NAME_KEYWORDS = {"test", "name", "item", "procedure", "service", "description",
                  "particular", "investigation", "examination", "profile"}

_SKIP_PATTERNS = re.compile(
    r"^(total|subtotal|page\s*\d|s\.no|sr\.?\s*no|sl\.?\s*no|"
    r"provider\s+name|catalogue|powered\s+by|price|rate|amount|code|"
    r"category|department|section|group|panel header)$",
    re.IGNORECASE,
)

_MIN_LEN = 2
_MAX_LEN = 200


def _is_valid_name(val: str) -> bool:
    val = val.strip()
    if len(val) < _MIN_LEN or len(val) > _MAX_LEN:
        return False
    if val.replace(" ", "").isdigit():
        return False
    if _SKIP_PATTERNS.match(val):
        return False
    return True


def _find_test_col_in_df(df: pd.DataFrame) -> Optional[int]:
    """Scan first 10 rows for a header that looks like a test-name column."""
    for row_i in range(min(10, len(df))):
        for col_i, val in enumerate(df.iloc[row_i]):
            if isinstance(val, str) and any(kw in val.lower() for kw in _NAME_KEYWORDS):
                return col_i
    return None


def _extract_from_excel(file_path: str) -> list[str]:
    xl = pd.ExcelFile(file_path)
    best: list[str] = []

    for sheet in xl.sheet_names:
        raw = pd.read_excel(file_path, sheet_name=sheet, header=None, dtype=str)
        if raw.empty:
            continue

        col_idx = _find_test_col_in_df(raw)
        if col_idx is None:
            col_idx = 0

        # Find the actual header row (first row where col_idx has a name-like value)
        header_row = 0
        for r in range(min(10, len(raw))):
            val = str(raw.iloc[r, col_idx] or "").strip()
            if val and any(kw in val.lower() for kw in _NAME_KEYWORDS):
                header_row = r + 1  # data starts below header
                break

        names = raw.iloc[header_row:, col_idx].dropna().astype(str).tolist()
        names = [n.strip() for n in names if _is_valid_name(n.strip())]

        if len(names) > len(best):
            best = names

    if not best:
        # Absolute fallback: first column of first sheet
        df = pd.read_excel(file_path, dtype=str)
        best = df.iloc[:, 0].dropna().astype(str).tolist()
        best = [n.strip() for n in best if _is_valid_name(n.strip())]

    return best


def _extract_from_pdf(file_path: str) -> list[str]:
    try:
        import pdfplumber
        names: list[str] = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Try tables first
                for table in (page.extract_tables() or []):
                    for row in table:
                        for cell in (row or []):
                            if cell and _is_valid_name(str(cell).strip()):
                                names.append(str(cell).strip())
                # Fall back to plain text
                if not names:
                    text = page.extract_text() or ""
                    for line in text.split("\n"):
                        line = line.strip()
                        if _is_valid_name(line):
                            names.append(line)
        return names
    except ImportError:
        pass

    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        names = []
        for page in doc:
            for line in (page.get_text() or "").split("\n"):
                line = line.strip()
                if _is_valid_name(line):
                    names.append(line)
        return names
    except ImportError:
        raise RuntimeError(
            "Install pdfplumber to process PDF files:  pip install pdfplumber"
        )


def _extract_from_docx(file_path: str) -> list[str]:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Install python-docx:  pip install python-docx")

    doc = Document(file_path)
    names: list[str] = []

    # Tables first (most catalogues are table-formatted)
    for table in doc.tables:
        # Detect which column is the test-name column from headers
        test_col = None
        if table.rows:
            header_cells = [c.text.strip() for c in table.rows[0].cells]
            for i, h in enumerate(header_cells):
                if any(kw in h.lower() for kw in _NAME_KEYWORDS):
                    test_col = i
                    break
            if test_col is None:
                test_col = 0

        for row_i, row in enumerate(table.rows):
            if row_i == 0:
                continue  # skip header
            if test_col < len(row.cells):
                text = row.cells[test_col].text.strip()
                if _is_valid_name(text):
                    names.append(text)

    # Paragraphs as fallback if no tables produced results
    if not names:
        for para in doc.paragraphs:
            text = para.text.strip()
            if _is_valid_name(text):
                names.append(text)

    return names


def _extract_from_csv(file_path: str) -> list[str]:
    df = pd.read_csv(file_path, dtype=str)
    for col in df.columns:
        if any(kw in col.lower() for kw in _NAME_KEYWORDS):
            return [n.strip() for n in df[col].dropna().astype(str).tolist() if _is_valid_name(n.strip())]
    return [n.strip() for n in df.iloc[:, 0].dropna().astype(str).tolist() if _is_valid_name(n.strip())]


def parse_file(file_path: str, jobs: dict, job_id: str) -> list[str]:
    ext = Path(file_path).suffix.lower()
    _progress(jobs, job_id, 10, f"Parsing {Path(file_path).name} …")

    if ext in (".xlsx", ".xls"):
        names = _extract_from_excel(file_path)
    elif ext == ".pdf":
        names = _extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        names = _extract_from_docx(file_path)
    elif ext == ".csv":
        names = _extract_from_csv(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Deduplicate preserving order
    seen: set[str] = set()
    unique: list[str] = []
    for n in names:
        key = n.lower().strip()
        if key not in seen:
            seen.add(key)
            unique.append(n)

    return unique


# ── match.py runner ───────────────────────────────────────────────────────────

def run_match_script(
    names: list[str],
    project_root: Path,
    job_id: str,
    jobs: dict,
) -> list[dict]:
    tmp_dir = Path(tempfile.gettempdir()) / "catalogue_jobs" / job_id
    tmp_dir.mkdir(parents=True, exist_ok=True)

    input_csv  = tmp_dir / "extracted_names.csv"
    output_csv = tmp_dir / "matched_results.csv"

    with open(input_csv, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Provider Test Name"])
        for n in names:
            writer.writerow([n])

    match_script = project_root / ".claude" / "skills" / "process-catalogue" / "scripts" / "match.py"
    master_file  = project_root / "refrences" / "Master.csv"

    _progress(jobs, job_id, 40, "Running match.py …")

    result = subprocess.run(
        [sys.executable, str(match_script),
         "--input",  str(input_csv),
         "--master", str(master_file),
         "--output", str(output_csv)],
        capture_output=True,
        text=True,
        timeout=180,
        cwd=str(project_root),
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"match.py exited with code {result.returncode}:\n{result.stderr[-2000:]}"
        )

    _progress(jobs, job_id, 70, "Reading match results …")
    results_df = pd.read_csv(output_csv)
    return results_df.to_dict("records")


# ── main job entry point ──────────────────────────────────────────────────────

def process_file_job(
    job_id: str,
    file_path: str,
    jobs: dict,
    project_root: Path,
) -> None:
    """Called by FastAPI BackgroundTasks — runs in a worker thread."""
    try:
        _progress(jobs, job_id, 5, "Checking dependencies …")
        _ensure_deps()

        # 1. Parse
        names = parse_file(file_path, jobs, job_id)
        if not names:
            raise ValueError("No test names could be extracted from the uploaded file.")

        _progress(jobs, job_id, 25, f"Extracted {len(names)} test names")

        # 2. Match
        results = run_match_script(names, project_root, job_id, jobs)

        _progress(jobs, job_id, 80, "Categorising results …")

        matched:   list[dict] = []
        unmatched: list[dict] = []
        skipped:   list[dict] = []

        for row in results:
            mt = str(row.get("Match Type", "UNMATCHED"))
            item = {
                "id":                     str(uuid.uuid4()),
                "provider_name":          str(row.get("Provider Test Name", "")),
                "catalogue_name":         str(row.get("Catalogue Test Name") or ""),
                "match_type":             mt,
                "confidence":             float(row.get("Confidence Score") or 0),
                "original_catalogue_name":str(row.get("Catalogue Test Name") or ""),
                "user_edited":            False,
            }
            if mt == "SKIPPED":
                skipped.append(item)
            elif mt == "UNMATCHED":
                item["suggested_name"] = ""
                unmatched.append(item)
            else:
                matched.append(item)

        stats = {
            "total":           len(names),
            "matched":         len(matched),
            "unmatched":       len(unmatched),
            "skipped":         len(skipped),
            "exact":           sum(1 for r in results if r.get("Match Type") == "exact"),
            "fuzzy":           sum(1 for r in results if r.get("Match Type") == "fuzzy"),
            "fuzzy_catalogue": sum(1 for r in results if r.get("Match Type") == "fuzzy-catalogue"),
        }

        jobs[job_id].update({
            "status":       "done",
            "progress":     100,
            "current_step": "Processing complete",
            "matched":      matched,
            "unmatched":    unmatched,
            "skipped":      skipped,
            "stats":        stats,
        })

    except Exception as exc:
        import traceback
        jobs[job_id].update({
            "status":       "error",
            "progress":     0,
            "current_step": f"Error: {exc}",
            "error":        str(exc),
            "traceback":    traceback.format_exc(),
        })


# ── Excel output generation ───────────────────────────────────────────────────

def _read_output_template_columns(project_root: Path) -> list[str]:
    """Read column names from Output_format.xlsx; fall back to safe defaults."""
    try:
        tpl = pd.read_excel(
            project_root / "refrences" / "Output_format.xlsx",
            nrows=0,
        )
        return list(tpl.columns)
    except Exception:
        return []


def generate_output_excel(
    job_id: str,
    matched_rows: list[dict],
    unmatched_rows: list[dict],
    filename: str,
    project_root: Path,
) -> str:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    # Same columns as the mapping results table UI — used for both Pathology and Radiology sheets
    TABLE_COLS = [
        ("provider_name",         "Raw Test Name"),
        ("price",                 "Price"),
        ("department",            "Dept"),
        ("catalogue_name",        "Standard Name"),
        ("provider_item_name",    "Provider Slug"),
        ("confidence",            "Confidence"),
        ("match_type",            "Status"),
        ("loinc_id",              "LOINC ID"),
        ("lab_requirement",       "Services Offered"),
        ("entity_type",           "Catalogue Type"),
        ("mrp",                   "Partner MRP (\u20b9)"),
        ("discounted_price",      "Discounted Price (\u20b9)"),
        ("display_mrp",           "Display MRP (\u20b9)"),
        ("collection_type",       "Lab Visit"),
        ("home_collection",       "Home Sample"),
        ("fasting_required",      "Fasting Required"),
        ("fasting_hours",         "Fasting Hours"),
        ("age_range",             "Age Range"),
        ("gender",                "Gender"),
        ("minimum_patient",       "Min. Patients"),
        ("report_tat",            "TAT (Hrs)"),
        ("alias",                 "Alias"),
        ("tags",                  "Tags"),
        ("prescription_required", "Rx Required"),
        ("precautions",           "Precautions"),
        ("description",           "Overview"),
    ]

    UNMATCHED_COLS = [
        ("provider_name",         "Raw Test Name"),
        ("catalogue_name",        "Suggested Standard Name"),
        ("department",            "Dept"),
        ("match_type",            "Status"),
        ("confidence",            "Confidence"),
    ]

    def _val(row: dict, key):
        if key is None:
            return ""
        if key == "confidence":
            v = row.get("confidence", 0)
            return f"{v:.0%}" if v else ""
        if key == "provider_item_name":
            cat = str(row.get("catalogue_name", "") or "").strip()
            return re.sub(r"[^a-z0-9]+", "-", cat.lower()).strip("-") if cat else ""
        return str(row.get(key, "") or "")

    def _write_sheet(ws, cols, rows, fill_hex, hdr_font):
        fill = PatternFill("solid", fgColor=fill_hex)
        for ci, (_, label) in enumerate(cols, 1):
            cell = ws.cell(row=1, column=ci, value=label)
            cell.fill = fill
            cell.font = hdr_font
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
        for ri, row in enumerate(rows, 2):
            for ci, (key, _) in enumerate(cols, 1):
                ws.cell(row=ri, column=ci, value=_val(row, key))
        ws.row_dimensions[1].height = 40
        for col_cells in ws.columns:
            width = max((len(str(c.value or "")) for c in col_cells), default=10)
            ws.column_dimensions[col_cells[0].column_letter].width = min(width + 4, 55)

    # Split matched rows by department
    path_rows  = [r for r in matched_rows if (r.get("department") or "").lower() == "pathology"]
    rad_rows   = [r for r in matched_rows if (r.get("department") or "").lower() == "radiology"]
    other_rows = [r for r in matched_rows if (r.get("department") or "").lower() not in ("pathology", "radiology")]
    path_rows += other_rows  # unknown dept falls into Pathology sheet

    # Build output path
    stem = re.sub(r"[^a-zA-Z0-9_\- ]", "_", Path(filename).stem)[:50].strip("_")
    output_dir = project_root / "output"
    output_dir.mkdir(exist_ok=True)
    base_name = f"{stem}_standardized_catalogue.xlsx"
    output_path = output_dir / base_name

    if output_path.exists():
        from datetime import datetime
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = output_dir / f"{stem}_standardized_catalogue_{ts}.xlsx"

    hdr_font = Font(color="FFFFFF", bold=True)
    wb = openpyxl.Workbook()

    ws_path = wb.active
    ws_path.title = "Pathology"
    _write_sheet(ws_path, TABLE_COLS, path_rows,  "1E3A5F", hdr_font)

    ws_rad = wb.create_sheet("Radiology")
    _write_sheet(ws_rad,  TABLE_COLS, rad_rows,   "14532D", hdr_font)

    ws_u = wb.create_sheet("UNMATCHED")
    _write_sheet(ws_u,    UNMATCHED_COLS, unmatched_rows, "7F1D1D", hdr_font)

    wb.save(str(output_path))
    return str(output_path)
